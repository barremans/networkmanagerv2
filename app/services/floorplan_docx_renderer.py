# =============================================================================
# Networkmap_Creator
# File:    app/services/floorplan_docx_renderer.py
# Role:    G-OPEN-8 — Grondplan export als Word-document (.docx)
#          Pure Python via python-docx — geen Node.js of externe runtime nodig.
#          Vereiste: pip install python-docx
# Version: 2.16.0
# Author:  Barremans
# Changes: 2.16.0 — Fix: from __future__ import annotations verplaatst
#                   naar regel 1 (SyntaxError bij import opgelost).
#                   Bedrijfsnaam toegevoegd aan paginaheader en ALL-header
#                   via get_company_for_site().
#          2.15.0 — Lege pagina's ALL-export opgelost:
#                   page_break_after=True bij PNG vervangen door False;
#                   één expliciete _add_page_break na PNG/sectietitel.
#                   Dubbele break (sectietitel-break + PNG-break) vermeden.
#          2.14.0 — FE-NEW-1: _resolve port-type: trace-stappen correct tonen
#                   inclusief wandpunt en eindapparaat.
#                   Back-poort: reversed() toegepast (zelfde logica als main_window).
#                   _add_trace: slice verhoogd van 6 naar 8 voor volledige trace.
#          2.13.0 — _safe_save() toegevoegd: PermissionError als Word het bestand
#                    al open heeft → leesbare foutmelding ipv crash
#          2.12.0 — header_distance vergroot zodat 20pt titel volledig zichtbaar is
#                    Datum blijft rechts op dezelfde regel via tab
#          2.11.0 — Header-titel naar Pt(20)
#          2.10.0 — Header-titel naar Pt(16)
#                    Gelijk aan sectietitel-font grondplannen
#                    Site-naam in header van Pt(11) naar Pt(13)
#          2.9.0 — Paginanummering toegevoegd rechtsonder: 'Pagina X van Y'
#                   Via Word-veldcodes PAGE + NUMPAGES in footer
#                   Geldt voor alle pagina's (single + ALL)
#          2.8.0 — Lege invultabel toegevoegd na kaartjes (single + ALL)
#                   Niet-gekoppelde SVG-punten als vooraf ingevulde ID-kolom
#                   5 extra lege rijen als buffer
#                   Kolommen: Wandpunt-Device ID | Port | Note | End-user
#          2.7.0 — PNG schaalt automatisch op basis van aspect-ratio
#                   MAX_H_CM=17.0cm bewaakt dat PNG niet over pagina loopt
#                   Afmetingen via PNG-header (struct) — geen Pillow nodig
#          2.6.0 — Lege pagina na PNG bij single export opgelost
#                   page_break_after=False bij single: sectie-wissel IS de page break
#                   page_break_after=True blijft correct voor ALL-export
#          2.5.0 — Lege pagina ALL opgelost, inhoudsopgave, sectietitel
#          2.4.0 — export_all_floorplans_docx() toegevoegd voor site-scope
#                   Elk grondplan krijgt eigen PNG + kaartjes als sectie
#                   Grondplan-naam als sectietitel tussen secties
#          2.3.0 — Kaartjes compact: ~6.5-7cm hoog → 2 per A4-landscape pagina
#                   Font 9pt in tabellen en trace, 11pt in badge
#                   Cel-padding gehalveerd voor compactheid
#                   Hoogte-schattingen bijgewerkt: outlet 6.8, ep 6.2, port 6.5cm
#          2.2.0 — Lege pagina verwijderd, overflow-logica gecorrigeerd
#          2.0.0 — Volledige herschrijving naar python-docx
#          1.0.0 — G-OPEN-8: initiële versie (Node.js)
#
# Installatie vereiste (eenmalig):
#   pip install python-docx
# =============================================================================

from __future__ import annotations

from app.helpers.settings_storage import get_all_sites, get_company_for_site

from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    pass

# Module-level imports (python-docx)
try:
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    # Wordt netjes afgevangen in export_floorplan_docx()
    pass


# ---------------------------------------------------------------------------
# Publieke entry point
# ---------------------------------------------------------------------------

def _safe_save(doc, filepath: str):
    """
    Sla het document op. Controleert eerst of het bestand beschikbaar is
    (niet vergrendeld door Word of een ander programma).
    Gooit PermissionError als het bestand in gebruik is.
    """
    import os
    # Test schrijftoegang door het bestand kort te openen
    if os.path.exists(filepath):
        try:
            with open(filepath, 'a'):
                pass
        except PermissionError:
            raise PermissionError(filepath)
    doc.save(filepath)


def export_floorplan_docx(
    floorplan: dict,
    site: dict,
    data: dict,
    filepath: str,
    png_path: str | None = None,
) -> tuple[bool, str]:
    """
    G-OPEN-8 — Exporteer gekoppelde grondplan-punten als Word-document.

    Parameters:
        floorplan : floorplan dict (id, name, mappings, ...)
        site      : site dict (id, name, ...)
        data      : volledige network_data dict
        filepath  : volledig uitvoerpad incl. .docx extensie
        png_path  : optioneel pad naar grondplan PNG (pagina 1 in Word)

    Returns:
        (True,  "")           bij succes
        (False, foutmelding)  bij fout
    """
    try:
        from docx import Document  # noqa — test import
    except ImportError:
        return False, (
            "python-docx niet geïnstalleerd.\n"
            "Voer uit: pip install python-docx"
        )

    try:
        doc = _build_document(floorplan, site, data, png_path=png_path)
        _safe_save(doc, filepath)
        return True, ""
    except PermissionError:
        return False, (
            f"Het bestand is al geopend in een ander programma (Word?).\n"
            f"Sluit het document eerst en probeer opnieuw:\n{filepath}"
        )
    except Exception:
        import traceback
        return False, traceback.format_exc()


def export_all_floorplans_docx(
    floorplans: list,
    site: dict,
    data: dict,
    filepath: str,
    png_paths: dict | None = None,
) -> tuple:
    """
    G-OPEN-8 — Exporteer alle grondplannen van een site als één Word-document.

    Parameters:
        floorplans : lijst van floorplan dicts
        site       : site dict
        data       : volledige network_data dict
        filepath   : uitvoerpad .docx
        png_paths  : dict {floorplan_id: png_pad} optioneel

    Returns:
        (True, "") bij succes, (False, foutmelding) bij fout
    """
    try:
        from docx import Document
    except ImportError:
        return False, "python-docx niet geïnstalleerd.\nVoer uit: pip install python-docx"

    if not floorplans:
        return False, "Geen grondplannen opgegeven."

    try:
        doc = _build_all_document(floorplans, site, data, png_paths or {})
        _safe_save(doc, filepath)
        return True, ""
    except PermissionError:
        return False, (
            f"Het bestand is al geopend in een ander programma (Word?).\n"
            f"Sluit het document eerst en probeer opnieuw:\n{filepath}"
        )
    except Exception:
        import traceback
        return False, traceback.format_exc()


# ---------------------------------------------------------------------------
# Kleuren en constanten
# ---------------------------------------------------------------------------

_C_ACCENT  = RGBColor(0x1F, 0x5C, 0x99)   # donkerblauw
_C_GROEN   = RGBColor(0x4C, 0xAF, 0x7D)   # wandpunt
_C_BLAUW   = RGBColor(0x21, 0x96, 0xF3)   # eindapparaat
_C_ORANJE  = RGBColor(0xFF, 0x70, 0x43)   # poort
_C_ZWART   = RGBColor(0x11, 0x11, 0x11)
_C_SUBTXT  = RGBColor(0x55, 0x55, 0x55)
_C_GRIJS   = RGBColor(0xF4, 0xF4, 0xF4)
_C_GRIJSDK = RGBColor(0xE0, 0xE0, 0xE0)
_C_WIT     = RGBColor(0xFF, 0xFF, 0xFF)

# Badge achtergronden per type
_BADGE_BG = {
    "outlet": RGBColor(0xE8, 0xF5, 0xEE),
    "ep":     RGBColor(0xE3, 0xF2, 0xFD),
    "port":   RGBColor(0xFB, 0xE9, 0xE7),
}
_TYPE_COLOR = {
    "outlet": _C_GROEN,
    "ep":     _C_BLAUW,
    "port":   _C_ORANJE,
}

def _rgb_hex(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


# ---------------------------------------------------------------------------
# Document opbouwen
# ---------------------------------------------------------------------------

# Geschatte kaartiehoogtes in cm (A4 landscape, 21cm - 3cm marges = 18cm bruikbaar)
_CARD_H_OUTLET = 6.8   # wandpunt + eindapparaat + trace (compact)
_CARD_H_EP     = 6.2   # eindapparaat + trace (compact)
_CARD_H_PORT   = 6.5   # device + poort + trace (compact)
_CARD_SEP      = 0.5   # ruimte tussen kaartjes
_PAGE_H_USABLE = 17.5  # bruikbare paginahoogte in cm (na header/marges)


def _card_height_cm(card_type: str) -> float:
    return {"outlet": _CARD_H_OUTLET, "ep": _CARD_H_EP}.get(card_type, _CARD_H_PORT)


def _build_document(floorplan: dict, site: dict, data: dict,
                    png_path: str | None = None):
    from docx import Document
    from docx.shared import Cm
    from docx.enum.section import WD_ORIENT

    doc = Document()

    # Standaard font instellen
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # ── Pagina 1: grondplan PNG ──────────────────────────────────────────
    sec0 = doc.sections[0]
    sec0.orientation   = WD_ORIENT.LANDSCAPE
    sec0.page_width    = Cm(29.7)
    sec0.page_height   = Cm(21.0)
    sec0.left_margin   = Cm(1.0)
    sec0.right_margin  = Cm(1.0)
    sec0.top_margin    = Cm(1.2)
    sec0.bottom_margin = Cm(1.0)

    _add_header(doc, floorplan, site, data)
    _add_page_numbers(sec0)

    if png_path and Path(png_path).exists():
        # Geen page_break_after — de sectie-paragraaf hieronder zorgt zelf
        # voor de pagina-wissel. page_break_after=True hier zou een lege
        # pagina veroorzaken (dubbele break).
        _add_floorplan_image(doc, png_path, page_break_after=False)
    else:
        p = doc.add_paragraph()
        p.add_run("(Grondplan afbeelding niet beschikbaar)").italic = True

    # ── Pagina 2+: kaartjes — sectie-wissel IS de page break ─────────────
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    new_sec_para = doc.add_paragraph()
    new_sec_para.paragraph_format.space_before = Pt(0)
    new_sec_para.paragraph_format.space_after  = Pt(0)
    # Sectie-eigenschappen aanpassen via XML
    pPr = new_sec_para._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    for tag, val in [('w:pgSz', {'w:w': '16838', 'w:h': '11906', 'w:orient': 'landscape'}),
                     ('w:pgMar', {'w:top': '851', 'w:right': '851',
                                  'w:bottom': '851', 'w:left': '851',
                                  'w:header': '708', 'w:footer': '708'})]:
        el = OxmlElement(tag)
        for k, v in val.items():
            el.set(qn(k), v)
        sectPr.append(el)
    pPr.append(sectPr)

    # Data verzamelen
    mappings   = floorplan.get("mappings", {})
    outlet_map = _build_outlet_map(data)
    device_map = {d["id"]: d for d in data.get("devices", [])}
    port_map   = {p["id"]: p for p in data.get("ports", [])}
    ep_map     = {e["id"]: e for e in data.get("endpoints", [])}

    from app.services import tracing as _tr
    from app.helpers.settings_storage import get_outlet_location_label

    items        = sorted(mappings.items())
    page_used_cm = 0.0   # hoeveel cm al gebruikt op huidige pagina

    for i, (svg_pt, mapped_val) in enumerate(items):
        entry  = _resolve(svg_pt, mapped_val, data,
                          outlet_map, device_map, port_map, ep_map,
                          _tr, get_outlet_location_label)
        card_h = _card_height_cm(entry["type"])

        if page_used_cm > 0:
            # Controleer of kaartje nog op huidige pagina past
            if page_used_cm + _CARD_SEP + card_h > _PAGE_H_USABLE:
                # Past niet meer — nieuwe pagina
                _add_page_break(doc)
                page_used_cm = 0.0
            else:
                # Past nog — alleen scheidingsruimte
                _spacer(doc, Cm(0.3))
                page_used_cm += _CARD_SEP

        _add_card(doc, entry)
        page_used_cm += card_h

    # ── Lege invultabel voor niet-gekoppelde SVG-punten ───────────────────
    from app.services import floorplan_svg_service as _fsvg
    from app.services import floorplan_service as _fps
    svg_path = _fps.get_svg_path(floorplan)
    if svg_path.exists():
        try:
            all_labels = set(_fsvg.detect_point_labels(svg_path))
            uncoupled  = sorted(all_labels - set(mappings.keys()))
        except Exception:
            uncoupled = []
    else:
        uncoupled = []

    if uncoupled:
        tbl_h = 1.5 + (len(uncoupled) + 5) * 0.8
        if page_used_cm > 0 and page_used_cm + 1.0 + tbl_h > _PAGE_H_USABLE:
            _add_page_break(doc)
        _add_empty_mapping_table(doc, uncoupled)

    return doc


def _add_floorplan_image(doc, png_path: str, page_break_after: bool = False):
    """
    Voeg grondplan PNG in.
    Schaalt automatisch zodat de afbeelding altijd binnen A4 landscape past,
    ongeacht de SVG-ratio. page_break_after=True voegt break toe in dezelfde run.
    """
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Beschikbare ruimte op A4 landscape na marges en header
    MAX_W_CM = 27.0
    MAX_H_CM = 17.0   # conservatief: 21cm - 2.2cm marges - 1.5cm header - 0.3cm buffer

    # Afmetingen uitlezen via python-docx interne helper (ondersteunt PNG + JPEG)
    # of via struct voor PNG, of JPEG SOF marker voor JPEG
    try:
        import struct
        with open(png_path, 'rb') as f:
            header = f.read(4)

        if header[:4] == b'\x89PNG':
            # PNG: breedte op bytes 16-19, hoogte op 20-23
            with open(png_path, 'rb') as f:
                f.seek(16)
                img_w = struct.unpack('>I', f.read(4))[0]
                img_h = struct.unpack('>I', f.read(4))[0]
            ratio = img_w / max(img_h, 1)

        elif header[:2] == b'\xff\xd8':
            # JPEG: zoek SOF0/SOF2 marker voor afmetingen
            with open(png_path, 'rb') as f:
                f.read(2)  # skip SOI
                ratio = 1.41  # default
                while True:
                    marker = f.read(2)
                    if len(marker) < 2:
                        break
                    length = struct.unpack('>H', f.read(2))[0]
                    if marker[1] in (0xC0, 0xC1, 0xC2):  # SOF marker
                        f.read(1)  # precision
                        img_h = struct.unpack('>H', f.read(2))[0]
                        img_w = struct.unpack('>H', f.read(2))[0]
                        ratio = img_w / max(img_h, 1)
                        break
                    f.seek(length - 2, 1)
        else:
            ratio = 1.41

    except Exception:
        ratio = 1.41   # A4 landscape als fallback

    # Bereken optimale breedte: max breedte én max hoogte respecteren
    w_from_max_w = MAX_W_CM
    w_from_max_h = MAX_H_CM * ratio
    use_w = min(w_from_max_w, w_from_max_h)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

    run = p.add_run()
    run.add_picture(png_path, width=Cm(use_w))

    if page_break_after:
        br_run = p.add_run()
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        br_run._r.append(br)


# ---------------------------------------------------------------------------
# Paginaheader
# ---------------------------------------------------------------------------

def _add_header(doc, floorplan: dict, site: dict, data: dict | None = None):
    import datetime
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    header  = doc.sections[0].header
    fp_name  = floorplan.get("name", "") or floorplan.get("outlet_location_key", "")
    site_nm  = site.get("name", "")
    # Bedrijfsnaam ophalen indien data beschikbaar
    company  = get_company_for_site(data, site.get("id", "")) if data else None
    co_name  = company.get("name", "") if company else ""
    datum    = datetime.date.today().strftime("%d/%m/%Y")

    # Vergroot header-marge zodat 20pt titel past
    doc.sections[0].header_distance = Pt(28)

    # Verwijder standaard lege paragraaf
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)

    # Regel 1: titel links, datum rechts via tab
    p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

    r1 = p.add_run(f"{fp_name}")
    r1.bold           = True
    r1.font.size      = Pt(20)
    r1.font.color.rgb = _C_ZWART

    if co_name or site_nm:
        parts = []
        if co_name:
            parts.append(co_name)
        if site_nm:
            parts.append(site_nm)
        r2 = p.add_run("   |   " + "  —  ".join(parts))
        r2.font.size      = Pt(13)
        r2.font.color.rgb = _C_SUBTXT

    # Datum rechts via tab op zelfde regel
    p.add_run("\t")
    rd = p.add_run(datum)
    rd.font.size      = Pt(10)
    rd.font.color.rgb = _C_SUBTXT

    # Tab stop uiterst rechts (A4 landscape = 15840 twips content breedte)
    pPr  = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab  = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '15840')
    tabs.append(tab)
    pPr.append(tabs)

    # Scheidingslijn
    p_sep = header.add_paragraph()
    pPr2  = p_sep._p.get_or_add_pPr()
    pBdr  = OxmlElement('w:pBdr')
    bot   = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '12')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), _rgb_hex(_C_ACCENT))
    pBdr.append(bot)
    jc = pPr2.find(qn('w:jc'))
    if jc is not None:
        jc.addnext(pBdr)
    else:
        pPr2.append(pBdr)
    p_sep.paragraph_format.space_before = Pt(2)
    p_sep.paragraph_format.space_after  = Pt(0)


def _add_border_bottom(paragraph, hex_color: str, size: int = 8):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(size))
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), hex_color)
    pBdr.append(bot)
    # pBdr moet na jc maar vóór rPr staan in pPr — invoegen op juiste positie
    rPr = pPr.find(qn('w:rPr'))
    if rPr is not None:
        pPr.insert(list(pPr).index(rPr), pBdr)
    else:
        pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Kaartje toevoegen
# ---------------------------------------------------------------------------

def _add_card(doc, entry: dict):
    card_type = entry["type"]
    color     = _TYPE_COLOR.get(card_type, _C_ZWART)
    bg        = _BADGE_BG.get(card_type, _C_GRIJS)

    # Badge-balk
    _add_badge(doc, entry["svg_pt"], entry["type_label"], entry["obj_name"], color, bg)

    # Inhoud per type
    if card_type == "outlet":
        _add_card_outlet(doc, entry, color)
    elif card_type == "ep":
        _add_card_ep(doc, entry, color)
    else:
        _add_card_port(doc, entry, color)

    # Trace
    _add_trace(doc, entry.get("trace_steps", []))

    # Kleine ruimte onderaan
    _spacer(doc, Pt(6))


def _add_badge(doc, svg_pt: str, type_label: str, obj_name: str,
               color: RGBColor, bg: RGBColor):
    from docx.shared import Cm
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    _shade_cell(cell, _rgb_hex(bg))
    _set_cell_margins(cell, top=60, bottom=60, left=120, right=120)
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    r1 = p.add_run(f"{svg_pt}   ")
    r1.bold      = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = _C_ZWART

    r2 = p.add_run(type_label)
    r2.bold      = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = color

    if obj_name:
        r3 = p.add_run(f"   |   {obj_name}")
        r3.bold      = True
        r3.font.size = Pt(11)
        r3.font.color.rgb = _C_ZWART

    _spacer(doc, Pt(2))


# ---------------------------------------------------------------------------
# Kaartje-types
# ---------------------------------------------------------------------------

def _add_card_outlet(doc, entry: dict, color: RGBColor):
    wo = entry.get("outlet") or {}
    ep = entry.get("endpoint")

    wo_rows = [
        ("Naam:",      wo.get("name", "")),
        ("Locatie:",   entry.get("outlet_location_label") or wo.get("location_description", "")),
        ("VLAN:",      wo.get("vlan", "")),
        ("Notities:",  wo.get("notes", "")),
    ]
    ep_rows = []
    if ep:
        ep_rows = [
            ("Naam:",       ep.get("name", "")),
            ("Type:",       ep.get("type", "")),
            ("IP adres:",   ep.get("ip", "")),
            ("MAC adres:",  ep.get("mac", "")),
            ("S/N:",        ep.get("serial", "")),
            ("Merk:",       ep.get("brand", "")),
            ("Model:",      ep.get("model", "")),
            ("Notities:",   ep.get("notes", "")),
        ]
    else:
        ep_rows = [("", "Geen eindapparaat")]

    _add_two_col_section(doc, "Wandpunt", wo_rows, "Eindapparaat", ep_rows, color)


def _add_card_ep(doc, entry: dict, color: RGBColor):
    ep = entry.get("endpoint") or {}
    left = [
        ("Naam:",      ep.get("name", "")),
        ("Type:",      ep.get("type", "")),
        ("IP adres:",  ep.get("ip", "")),
        ("MAC adres:", ep.get("mac", "")),
    ]
    right = [
        ("S/N:",       ep.get("serial", "")),
        ("Merk:",      ep.get("brand", "")),
        ("Model:",     ep.get("model", "")),
        ("Locatie:",   ep.get("location", "")),
        ("Notities:",  ep.get("notes", "")),
    ]
    _add_two_col_section(doc, "Eindapparaat", left, "Details", right, color)


def _add_card_port(doc, entry: dict, color: RGBColor):
    dev  = entry.get("device") or {}
    port = entry.get("port")   or {}
    wo   = entry.get("outlet")

    dev_rows = [
        ("Naam:",   dev.get("name", "")),
        ("Type:",   dev.get("type", "")),
        ("IP:",     dev.get("ip", "")),
        ("MAC:",    dev.get("mac", "")),
        ("Model:",  dev.get("model", "")),
        ("Rack:",   entry.get("rack_location", "")),
    ]
    port_rows = [
        ("Naam:",   port.get("name", "")),
        ("Kant:",   (port.get("side", "") or "").upper()),
    ]

    if wo:
        wo_rows = [
            ("Naam:",      wo.get("name", "")),
            ("Locatie:",   entry.get("outlet_location_label") or wo.get("location_description", "")),
            ("VLAN:",      wo.get("vlan", "")),
            ("Notities:",  wo.get("notes", "")),
        ]
        _add_two_col_section(doc, "Wandpunt", wo_rows, "Device", dev_rows, color)
        _spacer(doc, Pt(4))
        _add_section_title(doc, "Poort", color)
        _add_info_table(doc, port_rows)
    else:
        _add_two_col_section(doc, "Device", dev_rows, "Poort", port_rows, color)


# ---------------------------------------------------------------------------
# Trace-sectie
# ---------------------------------------------------------------------------

def _add_trace(doc, steps: list):
    _add_section_title(doc, "Trace", _C_ACCENT)
    if not steps:
        p = doc.add_paragraph()
        r = p.add_run("Geen trace beschikbaar")
        r.font.size   = Pt(10)
        r.font.color.rgb = _C_SUBTXT
        r.italic      = True
        return

    for step in steps[:8]:
        obj_type = step.get("obj_type", "")
        label    = step.get("label", "")
        prefix   = "->" if obj_type == "port" else (">" if obj_type == "endpoint" else ">>")
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.left_indent  = Cm(0.4)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        rp = p.add_run(f"{prefix}  {label}")
        rp.font.size      = Pt(9)
        rp.font.color.rgb = _C_ZWART

    if len(steps) > 8:
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(f"... (+{len(steps)-8} stappen)")
        r.font.size      = Pt(9)
        r.font.color.rgb = _C_SUBTXT
        r.italic         = True


# ---------------------------------------------------------------------------
# Tabel-helpers
# ---------------------------------------------------------------------------

def _add_section_title(doc, label: str, color: RGBColor):
    from docx.shared import Cm
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    _shade_cell(cell, _rgb_hex(_C_GRIJSDK))
    _set_cell_margins(cell, top=40, bottom=40, left=100, right=100)
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(label)
    r.bold           = True
    r.font.size      = Pt(9)
    r.font.color.rgb = color


def _add_info_table(doc, rows: list[tuple[str, str]]):
    """Eén kolom label/waarde tabel."""
    if not rows:
        return
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = 'Table Grid'

    # Kolombreedte: label 30%, waarde 70%
    from docx.shared import Cm
    for i, (label, value) in enumerate(rows):
        cells = tbl.rows[i].cells
        bg = _rgb_hex(_C_GRIJS) if i % 2 == 0 else "FFFFFF"

        # Label
        _shade_cell(cells[0], bg)
        _set_cell_margins(cells[0], top=30, bottom=30, left=100, right=40)
        cells[0].paragraphs[0]._element.getparent().remove(cells[0].paragraphs[0]._element)
        p0 = cells[0].add_paragraph()
        p0.paragraph_format.space_before = Pt(0)
        p0.paragraph_format.space_after  = Pt(0)
        r0 = p0.add_run(label)
        r0.font.size      = Pt(9)
        r0.font.color.rgb = _C_SUBTXT

        # Waarde
        _shade_cell(cells[1], bg)
        _set_cell_margins(cells[1], top=30, bottom=30, left=40, right=100)
        cells[1].paragraphs[0]._element.getparent().remove(cells[1].paragraphs[0]._element)
        p1 = cells[1].add_paragraph()
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after  = Pt(0)
        r1 = p1.add_run(str(value or "—"))
        r1.bold           = True
        r1.font.size      = Pt(9)
        r1.font.color.rgb = _C_ZWART

    # Kolombreedte instellen
    _set_col_widths(tbl, [Cm(5), Cm(21.7)])


def _add_two_col_section(doc, title1: str, rows1: list,
                          title2: str, rows2: list, color: RGBColor):
    """Twee secties naast elkaar in één tabel."""
    from docx.shared import Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    half = Cm(13.0)
    gap  = Cm(0.7)

    # Buitentabel: 1 rij, 3 kolommen (links | gap | rechts)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'

    cells = tbl.rows[0].cells

    # Geen randen op buitentabel
    for c in cells:
        _remove_cell_borders(c)

    # Kolombreedte
    _set_col_widths(tbl, [half, gap, half])

    def fill_section(cell, title, rows):
        _set_cell_margins(cell, top=0, bottom=0, left=0, right=0)
        # Verwijder lege standaard-paragraaf
        for p in cell.paragraphs:
            p._element.getparent().remove(p._element)
        # Sectietitel
        p_title = cell.add_paragraph()
        _shade_para_cell(p_title, _rgb_hex(_C_GRIJSDK))
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after  = Pt(0)
        r = p_title.add_run(title)
        r.bold           = True
        r.font.size      = Pt(9)
        r.font.color.rgb = color

        # Rijen
        for i, (label, value) in enumerate(rows):
            p_row = cell.add_paragraph()
            bg    = _rgb_hex(_C_GRIJS) if i % 2 == 0 else "FFFFFF"
            _shade_para_cell(p_row, bg)
            p_row.paragraph_format.space_before = Pt(0)
            p_row.paragraph_format.space_after  = Pt(0)

            rl = p_row.add_run(f"{label}  ")
            rl.font.size      = Pt(9)
            rl.font.color.rgb = _C_SUBTXT

            rv = p_row.add_run(str(value or "—"))
            rv.bold           = True
            rv.font.size      = Pt(9)
            rv.font.color.rgb = _C_ZWART

    fill_section(cells[0], title1, rows1)
    # Middelste cel leeg laten
    for p in cells[1].paragraphs:
        p._element.getparent().remove(p._element)
    cells[1].add_paragraph()
    fill_section(cells[2], title2, rows2)

    _spacer(doc, Pt(4))


# ---------------------------------------------------------------------------
# XML / opmaak helpers
# ---------------------------------------------------------------------------

def _shade_cell(cell, hex_color: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    shd   = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _shade_para_cell(para, hex_color: str):
    """Achtergrondkleur via alinea-opmaak (shading op pPr)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)
    # Kleine padding
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.left_indent  = Cm(0.2)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar  = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def _remove_cell_borders(cell):
    """Verwijder celranden via python-docx ingebouwde methode."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    # Verwijder bestaande tcBdr als aanwezig
    existing = tcPr.find(qn('w:tcBdr'))
    if existing is not None:
        tcPr.remove(existing)
    # Voeg tcBdr toe met correcte element-volgorde (voor w:shd)
    tcBdr = OxmlElement('w:tcBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tcBdr.append(el)
    # tcBdr hoort voor shd in tcPr
    shd = tcPr.find(qn('w:shd'))
    if shd is not None:
        shd.addprevious(tcBdr)
    else:
        tcPr.append(tcBdr)


def _set_col_widths(table, widths):
    from docx.oxml.ns import qn
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW  = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    from docx.oxml import OxmlElement
                    tcW = OxmlElement('w:tcW')
                    tcPr.append(tcW)
                tcW.set(qn('w:w'),    str(int(widths[i].twips)))
                tcW.set(qn('w:type'), 'dxa')


def _add_page_break(doc):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p  = doc.add_paragraph()
    r  = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._r.append(br)


def _spacer(doc, size):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = size
    p.paragraph_format.space_after  = Pt(0)


def _add_page_numbers(doc_section):
    """
    Voeg paginanummering toe rechtsonder in de footer: "Pagina X van Y"
    Werkt via Word veldcodes: PAGE en NUMPAGES.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    footer = doc_section.footer
    # Verwijder bestaande lege paragraaf
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

    def _field_run(field_type: str) -> object:
        """Maak een Word-veldcode run aan (PAGE of NUMPAGES)."""
        r = p.add_run()
        fld = OxmlElement('w:fldChar')
        fld.set(qn('w:fldCharType'), 'begin')
        r._r.append(fld)

        r2 = p.add_run()
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = f' {field_type} '
        r2._r.append(instr)

        r3 = p.add_run()
        fld2 = OxmlElement('w:fldChar')
        fld2.set(qn('w:fldCharType'), 'separate')
        r3._r.append(fld2)

        r4 = p.add_run('1')   # placeholder — Word vervangt dit
        r4.font.size      = Pt(9)
        r4.font.color.rgb = _C_SUBTXT

        r5 = p.add_run()
        fld3 = OxmlElement('w:fldChar')
        fld3.set(qn('w:fldCharType'), 'end')
        r5._r.append(fld3)
        return r4

    # "Pagina "
    r_pre = p.add_run("Pagina ")
    r_pre.font.size      = Pt(9)
    r_pre.font.color.rgb = _C_SUBTXT

    _field_run('PAGE')

    r_mid = p.add_run(" van ")
    r_mid.font.size      = Pt(9)
    r_mid.font.color.rgb = _C_SUBTXT

    _field_run('NUMPAGES')



# ---------------------------------------------------------------------------
# ALL-document helpers
# ---------------------------------------------------------------------------

def _build_all_document(floorplans: list, site: dict, data: dict,
                        png_paths: dict) -> object:
    """
    Bouw één document met alle grondplannen.
    Per grondplan (enkel die met koppelingen):
      - Blauwe sectietitel
      - Grondplan PNG (indien beschikbaar)
      - Kaartjes op pagina's (2 per pagina)
    """
    from docx import Document
    from docx.shared import Cm
    from docx.enum.section import WD_ORIENT
    from app.services import tracing as _tr
    from app.helpers.settings_storage import get_outlet_location_label

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    sec0 = doc.sections[0]
    sec0.orientation   = WD_ORIENT.LANDSCAPE
    sec0.page_width    = Cm(29.7)
    sec0.page_height   = Cm(21.0)
    sec0.left_margin   = Cm(1.0)
    sec0.right_margin  = Cm(1.0)
    sec0.top_margin    = Cm(1.2)
    sec0.bottom_margin = Cm(1.0)

    _add_all_header(doc, site, data)
    _add_page_numbers(sec0)

    # Inhoudsopgave
    fp_met_koppelingen = [fp for fp in floorplans if fp.get("mappings")]
    _add_toc(doc, fp_met_koppelingen, site)

    outlet_map = _build_outlet_map(data)
    device_map = {d["id"]: d for d in data.get("devices", [])}
    port_map   = {p["id"]: p for p in data.get("ports", [])}
    ep_map     = {e["id"]: e for e in data.get("endpoints", [])}

    first_fp = True
    for fp in floorplans:
        mappings = fp.get("mappings", {})
        if not mappings:
            continue  # Grondplannen zonder koppelingen overslaan

        fp_name = fp.get("name", "") or fp.get("outlet_location_key", fp.get("id", ""))

        # Sectietitel op nieuwe pagina.
        # De page break staat HIER (voor de sectietitel), niet na de PNG.
        # Zo ontstaat er nooit een dubbele break (= lege pagina):
        #   oud: page_break_after=True in PNG + _add_page_break volgende iteratie
        #   nieuw: één _add_page_break hier, PNG met page_break_after=False,
        #          daarna één _add_page_break voor de kaartjes.
        if not first_fp:
            _add_page_break(doc)
        first_fp = False

        # Sectietitel — prominente banner + paginanummer-ankerpunt
        _add_fp_section_title(doc, fp_name, site.get("name", ""), show_subtitle=True)

        # PNG — page_break_after=False: de break naar de kaartjespagina
        # volgt hieronder via één expliciete _add_page_break.
        png = png_paths.get(fp.get("id", ""))
        if png and Path(png).exists():
            _add_floorplan_image(doc, png, page_break_after=False)
        _add_page_break(doc)

        # Kaartjes
        items        = sorted(mappings.items())
        page_used_cm = 0.0

        for svg_pt, mapped_val in items:
            entry  = _resolve(svg_pt, mapped_val, data,
                              outlet_map, device_map, port_map, ep_map,
                              _tr, get_outlet_location_label)
            card_h = _card_height_cm(entry["type"])

            if page_used_cm > 0:
                if page_used_cm + _CARD_SEP + card_h > _PAGE_H_USABLE:
                    _add_page_break(doc)
                    page_used_cm = 0.0
                else:
                    _spacer(doc, Cm(0.3))
                    page_used_cm += _CARD_SEP

            _add_card(doc, entry)
            page_used_cm += card_h

        # Lege invultabel voor niet-gekoppelde punten van dit grondplan
        from app.services import floorplan_svg_service as _fsvg
        from app.services import floorplan_service as _fps
        svg_path = _fps.get_svg_path(fp)
        if svg_path.exists():
            try:
                all_labels = set(_fsvg.detect_point_labels(svg_path))
                uncoupled  = sorted(all_labels - set(mappings.keys()))
            except Exception:
                uncoupled = []
        else:
            uncoupled = []
        if uncoupled:
            tbl_h = 1.5 + (len(uncoupled) + 5) * 0.8
            if page_used_cm > 0 and page_used_cm + 1.0 + tbl_h > _PAGE_H_USABLE:
                _add_page_break(doc)
            _add_empty_mapping_table(doc, uncoupled)

    return doc


def _add_all_header(doc, site: dict, data: dict | None = None):
    """Header voor het ALL-document."""
    import datetime
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    header    = doc.sections[0].header
    site_name = site.get("name", "")
    company   = get_company_for_site(data, site.get("id", "")) if data else None
    co_name   = company.get("name", "") if company else ""
    header_label = "  —  ".join(p for p in [co_name, site_name] if p)
    datum     = datetime.date.today().strftime("%d/%m/%Y")

    for p in header.paragraphs:
        p._element.getparent().remove(p._element)

    p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    r1 = p.add_run(f"Alle grondplannen  |  {header_label}" if header_label
                   else "Alle grondplannen")
    r1.bold           = True
    r1.font.size      = Pt(12)
    r1.font.color.rgb = _C_ZWART

    p.add_run("\t")
    rd = p.add_run(datum)
    rd.font.size      = Pt(10)
    rd.font.color.rgb = _C_SUBTXT

    pPr  = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab  = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '15840')
    tabs.append(tab)
    pPr.append(tabs)

    p_sep = header.add_paragraph()
    pPr2  = p_sep._p.get_or_add_pPr()
    pBdr  = OxmlElement('w:pBdr')
    bot   = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '12')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), _rgb_hex(_C_ACCENT))
    pBdr.append(bot)
    jc = pPr2.find(qn('w:jc'))
    if jc is not None:
        jc.addnext(pBdr)
    else:
        pPr2.append(pBdr)
    p_sep.paragraph_format.space_before = Pt(0)
    p_sep.paragraph_format.space_after  = Pt(0)


def _add_empty_mapping_table(doc, uncoupled_labels: list):
    """
    Lege invultabel voor niet-gekoppelde SVG-punten.
    Kolommen: Wandpunt-Device ID | Port | Note | End-user
    Rijen: één per ongekoppeld punt + 5 lege extra rijen als buffer.
    """
    from docx.shared import Cm

    _spacer(doc, Cm(0.5))

    # Sectielabel
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(6)
    r = p_title.add_run("Niet-gekoppelde punten — invultabel")
    r.bold           = True
    r.font.size      = Pt(11)
    r.font.color.rgb = _C_ACCENT

    # Kolombreedtes in DXA (1cm = 567)
    COL_ID   = int(3.5 * 567)
    COL_PORT = int(3.5 * 567)
    COL_NOTE = int(9.5 * 567)
    COL_USER = int(9.5 * 567)
    widths   = [COL_ID, COL_PORT, COL_NOTE, COL_USER]
    hdrs     = ["WANDPUNT-DEVICE ID", "PORT", "NOTE", "END-USER"]

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'

    # Kolomkoppen
    for i, (hdr, w) in enumerate(zip(hdrs, widths)):
        cell = tbl.rows[0].cells[i]
        _shade_cell(cell, _rgb_hex(_C_GRIJSDK))
        _set_cell_margins(cell, top=60, bottom=60, left=100, right=60)
        cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(hdr)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _C_ZWART

    # Rijen: ongekoppelde punten + 5 lege buffer-rijen
    all_rows = uncoupled_labels + [""] * 5
    for i, label in enumerate(all_rows):
        row = tbl.add_row()
        bg  = _rgb_hex(_C_GRIJS) if i % 2 == 0 else "FFFFFF"
        for j, w in enumerate(widths):
            cell = row.cells[j]
            _shade_cell(cell, bg)
            _set_cell_margins(cell, top=70, bottom=70, left=100, right=60)
            cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            if j == 0 and label:
                r = p.add_run(label)
                r.font.size = Pt(9); r.font.color.rgb = _C_ZWART
            else:
                p.add_run("")  # lege invulcel

    _set_col_widths(tbl, [Cm(w / 567) for w in widths])
    _spacer(doc, Cm(0.3))


def _add_toc(doc, floorplans: list, site: dict):
    """Inhoudsopgave: lijst van grondplannen met korte samenvatting."""
    from docx.shared import Cm

    # Koptitel
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(8)
    r = p_title.add_run("Inhoudsopgave")
    r.bold           = True
    r.font.size      = Pt(13)
    r.font.color.rgb = _C_ACCENT

    if not floorplans:
        p = doc.add_paragraph()
        p.add_run("Geen grondplannen met koppelingen.").italic = True
        _add_page_break(doc)
        return

    # Tabel: grondplan | aantal koppelingen
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'

    # Kolomkoppen
    hdrs = ["Grondplan", "Koppelingen", "Types"]
    widths_cm = [14.0, 5.0, 8.7]
    for i, (hdr, w) in enumerate(zip(hdrs, widths_cm)):
        cell = tbl.rows[0].cells[i]
        _shade_cell(cell, _rgb_hex(_C_GRIJSDK))
        _set_cell_margins(cell, top=60, bottom=60, left=100, right=60)
        cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(hdr)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = _C_ZWART

    # Rijen per grondplan
    for i, fp in enumerate(floorplans):
        mappings = fp.get("mappings", {})
        fp_name  = fp.get("name", "") or fp.get("outlet_location_key", fp.get("id",""))

        # Tel types
        n_outlet = sum(1 for v in mappings.values() if not v.startswith(("ep:","port:")))
        n_ep     = sum(1 for v in mappings.values() if v.startswith("ep:"))
        n_port   = sum(1 for v in mappings.values() if v.startswith("port:"))
        types_str = "  ".join(filter(None, [
            f"{n_outlet}× wandpunt" if n_outlet else "",
            f"{n_ep}× eindapparaat" if n_ep else "",
            f"{n_port}× poort" if n_port else "",
        ]))

        row = tbl.add_row()
        bg  = _rgb_hex(_C_GRIJS) if i % 2 == 0 else "FFFFFF"

        vals = [fp_name, str(len(mappings)), types_str]
        for j, (val, w) in enumerate(zip(vals, widths_cm)):
            cell = row.cells[j]
            _shade_cell(cell, bg)
            _set_cell_margins(cell, top=50, bottom=50, left=100, right=60)
            cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            rv = p.add_run(val)
            rv.font.size = Pt(10)
            rv.font.color.rgb = _C_ZWART
            if j == 0: rv.bold = True

    # Kolombreedte instellen
    from docx.shared import Cm as _Cm
    _set_col_widths(tbl, [_Cm(w) for w in widths_cm])

    _spacer(doc, Pt(6))
    _add_page_break(doc)


def _add_fp_section_title(doc, fp_name: str, site_name: str,
                          show_subtitle: bool = False):
    """Blauwe sectietitel per grondplan in het ALL-document."""
    import datetime
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    _shade_cell(cell, _rgb_hex(_C_ACCENT))
    _set_cell_margins(cell, top=140, bottom=140, left=240, right=240)
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

    r1 = p.add_run(fp_name)
    r1.bold           = True
    r1.font.size      = Pt(16)
    r1.font.color.rgb = _C_WIT

    if site_name:
        r2 = p.add_run(f"   |   {site_name}")
        r2.font.size      = Pt(12)
        r2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)

    if show_subtitle:
        datum = datetime.date.today().strftime("%d/%m/%Y")
        r3 = p.add_run(f"   —   {datum}")
        r3.font.size      = Pt(10)
        r3.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)

    _spacer(doc, Pt(8))


# ---------------------------------------------------------------------------
# Data-resolvers
# ---------------------------------------------------------------------------

def _build_outlet_map(data: dict) -> dict:
    outlet_map = {}
    for s in get_all_sites(data):
        for r in s.get("rooms", []):
            for wo in r.get("wall_outlets", []):
                outlet_map[wo["id"]] = wo
    return outlet_map


def _resolve(svg_pt, mapped_val, data, outlet_map, device_map,
             port_map, ep_map, _tr, get_outlet_location_label) -> dict:

    if mapped_val.startswith("ep:"):
        ep_id = mapped_val[3:]
        ep    = ep_map.get(ep_id, {})
        conn  = next(
            (c for c in data.get("connections", [])
             if (c.get("to_type") == "endpoint" and c["to_id"] == ep_id) or
                (c.get("from_type") == "endpoint" and c["from_id"] == ep_id)),
            None,
        )
        steps = []
        if conn:
            pid   = conn["from_id"] if conn.get("to_type") == "endpoint" else conn["to_id"]
            steps = _tr.trace_from_port(data, pid)
        return {
            "svg_pt": svg_pt, "mapped_val": mapped_val, "type": "ep",
            "type_label": "Eindapparaat", "obj_name": ep.get("name", ""),
            "outlet": None, "outlet_location_label": None,
            "endpoint": ep, "device": None, "port": None, "rack_location": None,
            "trace_steps": _fmt_steps(steps),
        }

    if mapped_val.startswith("port:"):
        port_id = mapped_val[5:]
        port    = port_map.get(port_id, {})
        dev     = device_map.get(port.get("device_id", ""), {})
        rack    = _find_rack(dev.get("id", ""), data)
        wo      = next((o for o in outlet_map.values()
                        if o.get("port_id") == port_id), None)
        loc_lbl = None
        if wo:
            loc_key = wo.get("location_description", "") or ""
            loc_lbl = get_outlet_location_label(loc_key) if loc_key else None
        steps = _tr.trace_from_port(data, port_id)
        # 2.14.0 — FE-NEW-1: back-poort trace loopt van poort → PP_front → switch,
        # maar wandpunt + eindapparaat zitten langs de back-kant.
        # Zelfde reversed()-logica als main_window._on_port_clicked:
        # reversed() alleen als de eerste stap de startpoort zelf is.
        port_side = (port.get("side", "") or "").lower()
        first_is_start = (
            steps and steps[0].get("obj_type") == "port"
            and steps[0].get("obj_id") == port_id
        )
        if port_side == "back" and first_is_start:
            steps = list(reversed(steps))
        return {
            "svg_pt": svg_pt, "mapped_val": mapped_val, "type": "port",
            "type_label": "Poort", "obj_name": dev.get("name", ""),
            "outlet": _outlet_d(wo), "outlet_location_label": loc_lbl,
            "endpoint": None, "device": _device_d(dev), "port": _port_d(port),
            "rack_location": rack,
            "trace_steps": _fmt_steps(steps),
        }

    # Wandpunt
    wo  = outlet_map.get(mapped_val, {})
    ep  = ep_map.get(wo.get("endpoint_id", ""), None) if wo else None
    loc_key = (wo.get("location_description", "") or "") if wo else ""
    loc_lbl = get_outlet_location_label(loc_key) if loc_key else None
    steps   = _tr.trace_from_wall_outlet(data, mapped_val)
    return {
        "svg_pt": svg_pt, "mapped_val": mapped_val, "type": "outlet",
        "type_label": "Wandpunt", "obj_name": wo.get("name", "") if wo else "",
        "outlet": _outlet_d(wo), "outlet_location_label": loc_lbl,
        "endpoint": _ep_d(ep) if ep else None,
        "device": None, "port": None, "rack_location": None,
        "trace_steps": _fmt_steps(steps),
    }


def _outlet_d(wo):
    if not wo: return None
    return {"name": wo.get("name",""), "location_description": wo.get("location_description",""),
            "vlan": wo.get("vlan",""), "notes": wo.get("notes","") or wo.get("description","")}

def _ep_d(ep):
    if not ep: return None
    return {"name": ep.get("name",""), "type": ep.get("type",""),
            "ip": ep.get("ip",""), "mac": ep.get("mac",""),
            "serial": ep.get("serial","") or ep.get("serial_number",""),
            "brand": ep.get("brand","") or ep.get("manufacturer",""),
            "model": ep.get("model",""), "location": ep.get("location",""),
            "notes": ep.get("notes","") or ep.get("description","")}

def _device_d(dev):
    if not dev: return None
    return {"name": dev.get("name",""), "type": dev.get("type",""),
            "ip": dev.get("ip",""), "mac": dev.get("mac",""),
            "model": dev.get("model","")}

def _port_d(port):
    if not port: return None
    return {"name": port.get("name",""), "side": port.get("side","")}

def _find_rack(device_id, data):
    if not device_id: return ""
    for s in get_all_sites(data):
        for r in s.get("rooms",[]):
            for ra in r.get("racks",[]):
                for sl in ra.get("slots",[]):
                    if sl.get("device_id") == device_id:
                        return f"{r['name']} / {ra['name']}"
    return ""

def _fmt_steps(steps):
    return [{"obj_type": s.get("obj_type",""), "label": s.get("label","")}
            for s in steps]