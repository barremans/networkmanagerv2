# =============================================================================
# Networkmap_Creator
# File:    app/services/report_generator.py
# Role:    Word rapport generator (python-docx)
# Version: 2.12.0
# Author:  Barremans
# Changes: 2.12.0 -- REVIEW-AP stap 2 (service-API): publieke
#                  enumerate_action_items(data) -> alle actie-objecten met
#                  status ('open'/'approved') + approval-record, voor het
#                  reviewvenster. Bouwt index/maps intern. Eén bron met het
#                  rapport zodat tellingen altijd overeenkomen.
#          2.11.0 -- REVIEW-AP stap 1: actiepunten met manueel goed te keuren
#                  uitzonderingen. Nieuwe object-niveau enumerator
#                  _enumerate_action_objects() met stabiele sleutels
#                  "<check>:<object_id>" = ENIGE bron voor telling + (later) de
#                  reviewtool. _compute_action_items() telt nu vanuit de OPEN
#                  objecten (goedgekeurde afgetrokken via data["approvals"]).
#                  Risico-/onverbonden-detailtabellen sluiten goedgekeurde uit.
#                  Nieuwe sectie _add_approved_exceptions() ("Goedgekeurde
#                  uitzonderingen") per bedrijf. Drempelcheck VLAN-info blijft
#                  aggregaat. Switch-overzicht dup-IP-vlag blijft informatief.
#          2.10.0 -- F-RPT BugA: access point gold als 'niet verbonden / actiepunt'
#                  wanneer de uplink via een wandpunt door een patchpaneel-keten
#                  liep (geen directe endpoint↔poort verbinding). Nieuwe helper
#                  _resolve_ap_uplink() bepaalt nu de terminerende ACTIEVE poort
#                  via tracing.trace_from_wall_outlet (of directe verbinding).
#                  Gebruikt in _add_wifi_overview (kolom 'Verbonden switch/poort')
#                  en _compute_action_items ('zonder switch/poort'-telling).
#          2.9.0 -- RW-rapport: volledige bedrijfsgroepering. render_report_docx
#                  rendert nu per bedrijf een Heading 1-sectie (paginabreuk
#                  vooraf, behalve het eerste) met alle rubrieken gefilterd op
#                  dat bedrijf via _scope_data_to_company(); rubrieken als
#                  Heading 2 -> automatische Word-TOC (_add_toc) toont bedrijven
#                  en rubrieken. Glossarium/orphans/revisie eenmalig globaal
#                  achteraan. Dubbel-IP redundantiecheck (#-stacks) was reeds
#                  afgedekt en blijft ongewijzigd.
#          2.8.1 -- Fix: site-samenvatting telt verbindingen nu op beide zijden
#                  (gelijk aan detail/cover); cover-actiepunten en actieplan
#                  delen nu _compute_action_items() zodat tellingen altijd
#                  gelijk zijn; nieuwe sectie _add_orphans_section() voor
#                  niet-geplaatste devices/eindapparaten (enkel all-rapport)
#          2.8.0 -- F3: rapport per bedrijf. render_report_docx(company_id="")
#                  filtert data op één bedrijf via _scope_data_to_company();
#                  bedrijfsblok op cover (naam/adres/BTW/tel/mail/website);
#                  org_name uit bedrijfsnaam bij bedrijfsselectie
#          2.7.0 -- F1: get_all_sites() voor v2 JSON
#          2.6.1 — Dubbele-IP: ook .<digits>-suffix herkend als stack (SWITCH 9.1–9.3)
#                  _stack_group() als aparte functie, IP-adressen als naam uitgesloten
#          2.6.0 — Cover: compacter (titels kleiner, spacers gereduceerd, font 9pt)
#                  Actiepunten als rij in stats-tabel (geen losse paragrafen meer)
#                  Tabelkolommen aangepast: label 8cm / waarde 11cm
#          2.5.1 — Dubbele-IP: #-groepen uitgesloten van conflict
#          2.5.0 — Cover: actiepuntentelling op pagina 1 geplaatst (geen aparte pagina 2)
#                  Lege pagina 6 verwijderd: spacer na switch-overzicht verwijderd
#          2.5.1 — Dubbele-IP: #-groepen (redundantiegroep/stack) uitgesloten van conflict
#          2.4.0 — Pagina 2 fix: VLAN aandachtspunten op cover, spacing compact
#                  Pagina 6 fix: dup-IP als tabelrij (geen losse page)
#                  Cover: validatiestatus + issue-aantallen (pre-compute)
#                  Actieplan: SFP/UTP actiepunt, AP-datakwaliteit, MAC hernoemd
#                  Rackbezetting: slot.height, Tot./Bezet/Vrij U
#          2.0.0 — Volledige herschrijving
#                  Nieuwe layout: site-groepering, floorplan-stijl
#                  Per site: rack/devices/poorten, wandpunten, verbindingen,
#                  direct verbonden, VLAN
#                  Globaal: risico-WO, onverbonden WO, eindapparaten,
#                  device-info, switch-overzicht, uplinks, Wi-Fi, actieplan,
#                  begrippenlijst, revisiehistoriek
#                  tblHeader op alle tabellen (headers herhalen bij pagina-overgang)
#                  MAC-normalisatie, IP-deduplicatie, locatie-verrijking
#          1.3.0 — Direct endpoint: _conn_label herkent to_type=="endpoint"
#          1.0.0 — Initieel
# =============================================================================

from __future__ import annotations
import re as _re
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.helpers.settings_storage import (
    get_all_sites, get_all_companies, get_company_by_id,
)

# ---------------------------------------------------------------------------
# Kleurenpalet (floorplan-stijl)
# ---------------------------------------------------------------------------
_C_ACCENT  = RGBColor(0x1F, 0x5C, 0x99)   # Donkerblauw
_C_GROEN   = RGBColor(0x4C, 0xAF, 0x7D)   # Groen — wandpunten
_C_BLAUW   = RGBColor(0x21, 0x96, 0xF3)   # Blauw — eindapparaten
_C_ROOD    = RGBColor(0xC0, 0x39, 0x2B)   # Rood — risico / fout
_C_ORANJE  = RGBColor(0xE6, 0x7E, 0x22)   # Oranje — waarschuwing
_C_ZWART   = RGBColor(0x11, 0x11, 0x11)   # Near-black — hoofdtekst
_C_SUBTXT  = RGBColor(0x55, 0x55, 0x55)   # Grijs — labels
_C_GRIJS   = RGBColor(0xF4, 0xF4, 0xF4)   # Lichtgrijs — afwisselende rijen
_C_GRIJSDK = RGBColor(0xE0, 0xE0, 0xE0)   # Donkergrijs — subheaders
_C_WIT     = RGBColor(0xFF, 0xFF, 0xFF)

def _rgb_hex(c: RGBColor) -> str:
    return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"

# ---------------------------------------------------------------------------
# Leesbare labels voor location_description keys (user-configurable in app)
# ---------------------------------------------------------------------------
_LOC_LABELS: dict[str, str] = {
    "mic_g":                    "OB MIC",
    "mtoob_g":                  "OB MTO",
    "refterob_g":               "OB REFTER",
    "ateliergvk_g":             "OB ATELIER GVK",
    "thegatheringob_g":         "OB THE GATHERING",
    "concentrationcube1ob_g":   "OB CONCENTRATION CUBE 1",
    "concentrationcube2ob_g":   "OB CONCENTRATION CUBE 2",
    "inspirationstation1ob_g":  "OB INSPIRATION STATION 1",
    "hrob_g":                   "OB HR",
    "inkomob_g":                "OB INKOMHAL",
    "installatieob_g":          "OB INSTALLATIE EN ISOLATIE",
    "ob_verdiep1_gang":         "OB VERDIEP 1 GANG",
    "ob_compressor_g":          "OB COMPRESSOR LB",
    "ob_outside_g":             "OB OUTSIDE",
    "ob_productie":             "OB PRODUCTIE LB",
    "ob_mag_versp_g_l1":        "OB MAG VERSP L1",
    "ob_mag_versp_g_l2":        "OB MAG VERSP L2",
    "ob_mag_versp_g_l3":        "OB MAG VERSP L3",
    "verspaning_g":             "OB VERSPANING",
    "serverruimteob_g":         "OB SERVERLOKAAL",
    "it_marketing_bureau":      "OB STAFF",
    "containerb_a":             "CONTAINER B",
    "containerc_a":             "CONTAINER C",
    "containerd_a":             "CONTAINER D",
    "containerg_a":             "CONTAINER G",
    "refternb_g":               "REFTER NB",
    "bureelnb_g":               "BUREEL NB",
    "magazijn_g":               "MAGAZIJN NB",
    "production_a":             "PRODUCTIE",
    "magazijn_a":               "MAGAZIJN",
}

# Type mappings
_TYPE_MAP: dict[str, str] = {
    "switch":             "Switch",
    "patch_panel":        "Patchpanel",
    "patchpanel":         "Patchpanel",
    "cable_management":   "Kabelgoot",
    "server":             "Server",
    "router":             "Router",
    "firewall":           "Firewall",
    "ups":                "UPS",
    "nuc":                "NUC / Mini-PC",
    "nuc1":               "NUC / Mini-PC",
    "access_point":       "AP",
    "ap":                 "AP",
    "media_converter":    "Mediaconverter",
    "media_conv":         "Mediaconverter",
    "fiber_converter":    "Fiber connectie",
    "fiber":              "Fiber connectie",
    "mediaconverter":     "Mediaconverter",
    "nvr":                "NVR IPCAM",
    "smartlogger":        "Smartlogger",
    "pc":                 "PC",
    "laptop":             "Laptop",
    "sonos":              "Sonos server",
    "sonos_server":       "Sonos server",
    "gyron":              "Toegangscontroller",
    "distribution_plug":  "Verdeelstekker",
    "other":              "Ander",
    "ander":              "Ander",
    "verdeelstekker":     "Verdeelstekker",
}

_EP_TYPE_MAP: dict[str, str] = {
    "pc":                 "PC",
    "laptop":             "Laptop",
    "all_in_one":         "All-in-One",
    "access_point":       "Access Point",
    "printer":            "Printer",
    "plotter":            "Plotter",
    "ip_phone":           "IP Telefoon",
    "phone":              "Telefoon",
    "ip_camera":          "IP Camera",
    "camera":             "IP Camera",
    "conference_devices": "Vergaderapparaat",
    "docking_station":    "Docking Station",
    "nvr":                "NVR",
    "server":             "Server",
    "switch":             "Switch",
    "ot_machine":         "OT Machine",
    "ot":                 "OT Apparaat",
    "smartlogger":        "Smartlogger",
    "ups":                "UPS",
    "dali":               "DALI Controller",
    "nuc":                "NUC / Mini-PC",
    "sonos":              "Sonos",
    "display":            "Display",
    "other":              "Ander",
}

_CABLE_LABELS: dict[str, str] = {
    "utp_cat5e": "UTP Cat5e",
    "utp_cat6":  "UTP Cat6",
    "utp_cat6a": "UTP Cat6a",
    "sfp_fiber": "SFP Fiber",
    "sfp_dac":   "SFP DAC",
    "other":     "Anders",
}

# ===========================================================================
# CELL / TABLE HELPERS
# ===========================================================================

def _shade(cell, hex_color: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    ex = tcPr.find(qn("w:shd"))
    if ex is not None:
        tcPr.remove(ex)
    tcPr.append(shd)


def _margins(cell, top: int = 80, bottom: int = 80,
             left: int = 120, right: int = 120) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar  = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    ex = tcPr.find(qn("w:tcMar"))
    if ex is not None:
        tcPr.remove(ex)
    tcPr.append(mar)


def _no_borders(cell) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    bdr  = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        bdr.append(el)
    ex = tcPr.find(qn("w:tcBorders"))
    if ex is not None:
        tcPr.remove(ex)
    tcPr.append(bdr)


def _col_widths(table, widths: list) -> None:
    for i, col in enumerate(table.columns):
        if i < len(widths):
            col.width = widths[i]
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            if j < len(widths):
                cell.width = widths[j]


def _cant_split(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    cs   = OxmlElement("w:cantSplit")
    cs.set(qn("w:val"), "1")
    trPr.append(cs)


def _tbl_header(row) -> None:
    """Markeer rij als herhaalbare tabelheader bij pagina-overgang."""
    trPr = row._tr.get_or_add_trPr()
    th   = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "1")
    trPr.append(th)


def _clear_p(cell) -> None:
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)


def _cell_p(cell, text: str, bold: bool = False, size: int = 9,
            color: RGBColor | None = None,
            align: WD_ALIGN_PARAGRAPH | None = None):
    _clear_p(cell)
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    return r


def _add_border_bottom(paragraph, color: RGBColor, size: int = 8) -> None:
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(size))
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), _rgb_hex(color))
    pBdr.append(bot)
    pPr.append(pBdr)


def _spacer(doc, cm: float = 0.3) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Cm(cm)
    p.paragraph_format.space_after  = Pt(0)


_SUPPRESS_NEXT_BREAK = False


def _add_page_break(doc) -> None:
    global _SUPPRESS_NEXT_BREAK
    if _SUPPRESS_NEXT_BREAK:
        _SUPPRESS_NEXT_BREAK = False
        return
    p  = doc.add_paragraph()
    r  = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r._r.append(br)


# ---------------------------------------------------------------------------
# Volledige-breedte cel helper
# ---------------------------------------------------------------------------
def _full_cell(doc, bg_hex: str, top: int = 80, bottom: int = 80):
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _shade(cell, bg_hex)
    _no_borders(cell)
    _margins(cell, top=top, bottom=bottom, left=200, right=200)
    _col_widths(tbl, [Cm(27.7)])
    return cell


# ===========================================================================
# DATA HELPERS
# ===========================================================================

def _wo_location(wo: dict) -> str:
    """Leesbare locatienaam: notes → _LOC_LABELS → clean fallback."""
    notes = (wo.get("notes") or "").strip()
    if notes:
        return notes
    loc = (wo.get("location_description") or "").strip()
    if not loc:
        return "—"
    if loc in _LOC_LABELS:
        return _LOC_LABELS[loc]
    clean = _re.sub(r"_(g|a|l\d+)$", "", loc)
    return clean.replace("_", " ").upper()


def _normalize_mac(mac: str) -> str:
    """Normaliseer MAC naar AA:BB:CC:DD:EE:FF."""
    if not mac:
        return "—"
    clean = _re.sub(r"[^0-9a-fA-F]", "", mac)
    if len(clean) == 12:
        return ":".join(clean[i:i+2].upper() for i in range(0, 12, 2))
    return mac.upper()


def _normalize_ip(ip: str) -> str:
    """Strip URL-prefix, geef enkel IP terug."""
    if not ip:
        return "—"
    ip = ip.strip()
    ip = _re.sub(r"^https?://", "", ip).rstrip("/")
    m  = _re.search(r"\d+\.\d+\.\d+\.\d+", ip)
    return m.group() if m else ip or "—"


def _dev_loc_str(idx: dict, dev_id: str) -> str:
    """Geeft ' (RUIMTE / RACK / U31)' voor een device_id."""
    loc = idx.get("loc", {}).get(dev_id)
    if not loc:
        return ""
    return f" ({loc['room']} / {loc['rack']} / U{loc['slot']})"


def _conn_label(data: dict, idx: dict, port_id: str) -> str:
    """Verbindingslabel voor een poort, incl. locatie-verrijking."""
    for conn in data.get("connections", []):
        other_id = other_type = None
        if conn.get("from_id") == port_id:
            other_id, other_type = conn.get("to_id"), conn.get("to_type")
        elif conn.get("to_id") == port_id:
            other_id, other_type = conn.get("from_id"), conn.get("from_type")
        if not other_id:
            continue

        if other_type in ("port", None):
            p2  = idx["port"].get(other_id)
            dev = idx["dev"].get(p2["device_id"]) if p2 else None
            if p2 and dev:
                side    = "F" if p2["side"] == "front" else "B"
                loc_str = _dev_loc_str(idx, dev["id"])
                return f"{dev['name']} / {p2['name']} ({side}){loc_str}"

        elif other_type == "wall_outlet":
            wo = idx["wo"].get(other_id)
            if wo:
                parts = [f"🌐 {wo.get('name', other_id)}"]
                loc   = _wo_location(wo)
                if loc and loc != "—":
                    parts.append(f"  ·  {loc}")
                ep = idx["ep"].get(wo.get("endpoint_id", ""))
                if ep:
                    parts.append(f"  —  {ep['name']}")
                return "".join(parts)

        elif other_type == "endpoint":
            ep = idx["ep"].get(other_id)
            if ep:
                return f"🖥 {ep.get('name', other_id)}"
    return "—"


def _resolve_side(data: dict, idx: dict, obj_id: str, obj_type: str) -> str:
    """Verbindingsuiteinde label (voor verbindingentabel), incl. locatie."""
    if not obj_id:
        return "—"
    if obj_type in ("port", None, ""):
        p   = idx["port"].get(obj_id)
        dev = idx["dev"].get(p["device_id"]) if p else None
        if p and dev:
            side    = "F" if p["side"] == "front" else "B"
            loc_str = _dev_loc_str(idx, dev["id"])
            return f"{dev['name']} / {p['name']} ({side}){loc_str}"
    elif obj_type == "wall_outlet":
        wo = idx["wo"].get(obj_id)
        if wo:
            wo_name = wo.get("name", obj_id)
            loc     = _wo_location(wo)
            return f"🌐 {wo_name}  ·  {loc}" if loc != "—" else f"🌐 {wo_name}"
    elif obj_type == "endpoint":
        ep = idx["ep"].get(obj_id)
        if ep:
            return f"🖥 {ep.get('name', obj_id)}"
    return obj_id


def _is_site_port(obj_id: str, obj_type: str,
                  idx: dict, site_dev_ids: set) -> bool:
    if obj_type not in ("port", None, ""):
        return False
    p = idx["port"].get(obj_id)
    return p is not None and p.get("device_id") in site_dev_ids


def _loc_group_label(loc_key: str) -> str:
    """Leesbaar label voor een locatie-groep sleutel."""
    if loc_key in _LOC_LABELS:
        return _LOC_LABELS[loc_key]
    if not loc_key:
        return "—"
    clean = _re.sub(r"_(g|a|l\d+)$", "", loc_key)
    return clean.replace("_", " ").upper()


# ===========================================================================
# INDEX BUILDER
# ===========================================================================

def _build_index(data: dict) -> dict:
    """Bouw snelle opzoek-index van alle objecten + locatie-index."""
    idx: dict = {
        "dev":  {d["id"]: d for d in data.get("devices", [])},
        "port": {p["id"]: p for p in data.get("ports",   [])},
        "ep":   {e["id"]: e for e in data.get("endpoints", [])},
        "wo":   {},
        "loc":  {},   # device_id → {site, room, rack, slot}
    }
    for site in get_all_sites(data):
        for room in site.get("rooms", []):
            for wo in room.get("wall_outlets", []):
                idx["wo"][wo["id"]] = wo
            for rack in room.get("racks", []):
                total_u = rack.get("total_units", 42)
                for slot in rack.get("slots", []):
                    dev_id = slot.get("device_id")
                    if dev_id:
                        u_start  = slot.get("u_start")
                        slot_lbl = str(total_u - u_start + 1) if u_start is not None else "?"
                        idx["loc"][dev_id] = {
                            "site": site["name"],
                            "room": room["name"],
                            "rack": rack["name"],
                            "slot": slot_lbl,
                        }
    return idx


def _is_patchpanel_type(dev_type: str) -> bool:
    """True voor passieve patchpanelen (beide schrijfwijzen komen voor)."""
    return dev_type in ("patch_panel", "patchpanel")


def _resolve_ap_uplink(data: dict, idx: dict, ap_id: str) -> str | None:
    """
    Bepaalt de terminerende ACTIEVE switch-/netwerkpoort voor een access point.

    Een AP kan op twee manieren verbonden zijn:
      1. directe endpoint->poort verbinding;
      2. via een wandpunt (wall_outlet.endpoint_id) dat door een patchpaneel-
         keten naar een switchpoort gepatcht is.

    Geeft het port_id van de LAATSTE niet-patchpaneel (actieve) poort in de
    trace terug, of None als de AP nergens op een actief toestel uitkomt.
    """
    # 1. directe endpoint<->poort verbinding
    for conn in data.get("connections", []):
        if (conn.get("to_type") == "endpoint" and conn.get("to_id") == ap_id
                and conn.get("from_type") == "port"):
            return conn.get("from_id")
        if (conn.get("from_type") == "endpoint" and conn.get("from_id") == ap_id
                and conn.get("to_type") == "port"):
            return conn.get("to_id")

    # 2. via gekoppeld wandpunt -> trace door de patchpaneel-keten
    outlet_id = None
    for site in get_all_sites(data):
        for room in site.get("rooms", []):
            for wo in room.get("wall_outlets", []):
                if wo.get("endpoint_id") == ap_id:
                    outlet_id = wo["id"]
                    break
            if outlet_id:
                break
        if outlet_id:
            break
    if not outlet_id:
        return None

    from app.services import tracing
    steps = tracing.trace_from_wall_outlet(data, outlet_id)
    last_active_port = None
    for st in steps:
        if st.get("obj_type") != "port":
            continue
        port = idx["port"].get(st.get("obj_id", ""))
        dev  = idx["dev"].get(port.get("device_id", "")) if port else None
        if dev and not _is_patchpanel_type(dev.get("type", "")):
            last_active_port = st.get("obj_id")
    return last_active_port


# ===========================================================================
# ACTIEPUNTEN — object-niveau enumerator + goedkeuringen (REVIEW-AP)
# ===========================================================================

def _action_key(check: str, obj_id: str) -> str:
    """Stabiele sleutel voor een actie-object: '<check>:<object_id>'."""
    return f"{check}:{obj_id}"


def _approved_keys(data: dict) -> dict:
    """Map van goedgekeurde actie-sleutels -> approval-record (status 'approved')."""
    out: dict = {}
    for a in data.get("approvals", []):
        if a.get("status") == "approved" and a.get("key"):
            out[a["key"]] = a
    return out


def _enumerate_action_objects(data: dict, idx: dict,
                              risk_by_site: dict, unwired_by_site: dict,
                              dup_ips: set) -> list:
    """
    Object-niveau actiepunten met stabiele sleutel. ENIGE bron voor zowel de
    telling (_compute_action_items) als de reviewtool, zodat aantallen altijd
    overeenkomen. Aggregaat-/drempelchecks (poorten zonder VLAN) zitten hier
    NIET in — die blijven categorie-niveau in _compute_action_items.

    Elk object: dict met key, check, category, object_type, object_id, label,
    detail, priority (+ optioneel _name / _desc / _mac voor de telling-teksten).
    """
    objs: list = []

    def add(check, otype, oid, label, detail, prio, **extra):
        o = {
            "key":         _action_key(check, oid),
            "check":       check,
            "category":    _ACTION_CATEGORY.get(check, check),
            "object_type": otype,
            "object_id":   oid,
            "label":       label,
            "detail":      detail,
            "priority":    prio,
        }
        o.update(extra)
        objs.append(o)

    # Risico wandpunten — verbonden zonder eindapparaat
    for sn, rows in risk_by_site.items():
        for (wo, rn, port, dev) in rows:
            add("risk_outlet", "wall_outlet", wo["id"],
                f"{wo.get('name','?')}  ·  {rn}",
                "Verbonden aan netwerk zonder eindapparaat", "Hoog")

    # Onverbonden wandpunten — geen kabelverbinding
    for sn, rows in unwired_by_site.items():
        for (wo, rn) in rows:
            add("unwired_outlet", "wall_outlet", wo["id"],
                f"{wo.get('name','?')}  ·  {rn}",
                "Geen kabelverbinding geregistreerd", "Middel")

    # Dubbele IP-adressen
    for ip in sorted(dup_ips):
        add("dup_ip", "ip", ip, ip, "Meerdere devices delen dit IP-adres", "Hoog")

    # Eindapparaten zonder sitekoppeling
    ep_site_map = _build_ep_site_map(data, idx)
    for ep in data.get("endpoints", []):
        if ep["id"] not in ep_site_map:
            add("unknown_ep", "endpoint", ep["id"], ep.get("name", "?"),
                "Niet verbonden via poort of wandpunt", "Laag")

    # Wi-Fi access points — datakwaliteit (per AP, per subtype)
    aps = [ep for ep in data.get("endpoints", []) if ep.get("type") == "access_point"]
    for ep in aps:
        if not ep.get("ip"):
            add("ap_no_ip", "endpoint", ep["id"], ep.get("name", "?"),
                "Access point zonder IP-adres", "Middel")
        if _resolve_ap_uplink(data, idx, ep["id"]) is None:
            add("ap_no_uplink", "endpoint", ep["id"], ep.get("name", "?"),
                "Access point komt niet op een actieve switch uit", "Middel")
    _seen_mac: dict = {}
    for ep in aps:
        m = ep.get("mac", "")
        if m:
            _seen_mac.setdefault(m, []).append(ep)
    for m, eps in _seen_mac.items():
        if len(eps) > 1:
            for ep in eps:
                add("ap_dup_mac", "endpoint", ep["id"], ep.get("name", "?"),
                    f"Dubbele identiteit (MAC {_normalize_mac(m)})", "Middel", _mac=m)

    # MAC-adressen in brondata niet genormaliseerd
    for dv in data.get("devices", []):
        mac = dv.get("mac", "")
        if mac and ":" not in mac and "-" not in mac:
            add("bad_mac", "device", dv["id"], dv.get("name", "?"),
                f"MAC zonder scheidingstekens: {mac}", "Laag", _name=dv.get("name", "?"))

    # Switches zonder merk/model
    for dv in data.get("devices", []):
        if dv.get("type") == "switch" and not dv.get("brand") and not dv.get("model"):
            add("no_brand_model", "device", dv["id"], dv.get("name", "?"),
                "Switch zonder merk/model", "Laag", _name=dv.get("name", "?"))

    # SFP-uplinks met mogelijk verkeerd kabeltype (per verbinding)
    _sw_ids = {dv["id"] for dv in data.get("devices", []) if dv.get("type") == "switch"}
    for c in data.get("connections", []):
        fp = idx.get("port", {}).get(c.get("from_id", ""), {})
        tp = idx.get("port", {}).get(c.get("to_id", ""), {})
        cb = (c.get("cable_type", "") or "")
        if ("SFP" in (fp.get("name", "") + tp.get("name", "")).upper()
                and any(kw in cb.upper() for kw in ("UTP", "CAT"))):
            fd = idx.get("dev", {}).get(fp.get("device_id", ""), {})
            td = idx.get("dev", {}).get(tp.get("device_id", ""), {})
            if fd.get("id") in _sw_ids and td.get("id") in _sw_ids:
                desc = (f"{fd.get('name','?')}/{fp.get('name','?')} -> "
                        f"{td.get('name','?')}/{tp.get('name','?')}")
                add("sfp_utp", "connection", c.get("id", ""), desc,
                    "SFP-uplink met UTP/CAT-kabeltype", "Middel", _desc=desc)

    return objs


_ACTION_CATEGORY = {
    "risk_outlet":     "Risico wandpunt",
    "unwired_outlet":  "Onverbonden wandpunt",
    "dup_ip":          "Dubbel IP-adres",
    "unknown_ep":      "Eindapparaat zonder sitekoppeling",
    "ap_no_ip":        "Wi-Fi datakwaliteit",
    "ap_no_uplink":    "Wi-Fi datakwaliteit",
    "ap_dup_mac":      "Wi-Fi datakwaliteit",
    "bad_mac":         "MAC niet genormaliseerd",
    "no_brand_model":  "Switch zonder merk/model",
    "sfp_utp":         "SFP-uplink kabeltype",
}


def _strip_approved_outlets(risk_by_site: dict, unwired_by_site: dict,
                            approved: dict) -> tuple[dict, dict]:
    """Verwijder goedgekeurde wandpunten uit de risico-/onverbonden-maps zodat
    detailtabellen én tellingen consistent blijven met het actieplan."""
    r2 = {}
    for sn, rows in risk_by_site.items():
        kept = [t for t in rows if _action_key("risk_outlet", t[0]["id"]) not in approved]
        if kept:
            r2[sn] = kept
    u2 = {}
    for sn, rows in unwired_by_site.items():
        kept = [t for t in rows if _action_key("unwired_outlet", t[0]["id"]) not in approved]
        if kept:
            u2[sn] = kept
    return r2, u2


def _collect_approved_objects(data: dict, idx: dict) -> list:
    """Alle goedgekeurde actie-objecten (volledige maps), met approval-record."""
    risk, unw = _build_wo_risk_maps(data, idx)
    dup       = _build_duplicate_ip_set(data)
    objs      = _enumerate_action_objects(data, idx, risk, unw, dup)
    approved  = _approved_keys(data)
    out = []
    for o in objs:
        ap = approved.get(o["key"])
        if ap:
            m = dict(o)
            m["_approval"] = ap
            out.append(m)
    return out


def enumerate_action_items(data: dict) -> list:
    """
    Publieke API voor de reviewtool. Geeft ALLE actie-objecten terug (open +
    goedgekeurd), elk met:
      key, check, category, object_type, object_id, label, detail, priority,
      status ('open' | 'approved'), approval (record of {}).
    Bouwt index en maps intern; identiek aan wat het rapport telt (één bron).
    """
    idx       = _build_index(data)
    risk, unw = _build_wo_risk_maps(data, idx)
    dup       = _build_duplicate_ip_set(data)
    objs      = _enumerate_action_objects(data, idx, risk, unw, dup)
    approved  = _approved_keys(data)
    for o in objs:
        ap = approved.get(o["key"])
        o["status"]   = "approved" if ap else "open"
        o["approval"] = ap or {}
    return objs


# ===========================================================================
# PRE-COMPUTE HELPERS
# ===========================================================================

def _build_ep_site_map(data: dict, idx: dict) -> dict:
    """ep_id → site_name"""
    ep_site: dict = {}
    port_idx = idx["port"]
    for site in get_all_sites(data):
        site_dev_ids = {
            slot.get("device_id")
            for room in site.get("rooms", [])
            for rack in room.get("racks", [])
            for slot in rack.get("slots", [])
            if slot.get("device_id")
        }
        for room in site.get("rooms", []):
            for wo in room.get("wall_outlets", []):
                ep_id = wo.get("endpoint_id")
                if ep_id:
                    ep_site[ep_id] = site["name"]
        for conn in data.get("connections", []):
            if conn.get("to_type") == "endpoint":
                p = port_idx.get(conn["from_id"])
                if p and p.get("device_id") in site_dev_ids:
                    ep_site[conn["to_id"]] = site["name"]
            if conn.get("from_type") == "endpoint":
                p = port_idx.get(conn["to_id"])
                if p and p.get("device_id") in site_dev_ids:
                    ep_site[conn["from_id"]] = site["name"]
    return ep_site


def _build_wo_risk_maps(data: dict, idx: dict) -> tuple[dict, dict]:
    """
    risk_by_site    — site → [(wo, room, port, dev)]  verbonden, geen EP
    unwired_by_site — site → [(wo, room)]             geen kabelverbinding
    """
    port_idx = idx["port"]
    dev_idx  = idx["dev"]
    wo_connected: dict = {}
    for conn in data.get("connections", []):
        for side, other in (("from", "to"), ("to", "from")):
            if conn.get(f"{side}_type") == "wall_outlet":
                wo_id      = conn.get(f"{side}_id")
                other_id   = conn.get(f"{other}_id")
                other_type = conn.get(f"{other}_type")
                if other_id and other_type in ("port", None):
                    p   = port_idx.get(other_id)
                    dev = dev_idx.get(p["device_id"]) if p else None
                    wo_connected[wo_id] = (p, dev)

    risk_by_site: dict    = {}
    unwired_by_site: dict = {}
    for site in get_all_sites(data):
        sn = site["name"]
        for room in site.get("rooms", []):
            rn = room["name"]
            for wo in room.get("wall_outlets", []):
                wid = wo["id"]
                if wid in wo_connected:
                    if not wo.get("endpoint_id"):
                        p, dev = wo_connected[wid]
                        risk_by_site.setdefault(sn, []).append((wo, rn, p, dev))
                else:
                    unwired_by_site.setdefault(sn, []).append((wo, rn))
    return risk_by_site, unwired_by_site


def _stack_group(name: str) -> str:
    """
    Leid de redundantiegroepsnaam af uit een device-naam.
    Twee patronen worden herkend:
      #<digits>  suffix  —  SW01#1, SW01#2       →  'SW01'
      .<digits>  suffix  —  SWITCH 9.1, 9.2     →  'SWITCH 9'
    Vereiste voor het punt-patroon: het prefix moet minstens één
    niet-numeriek teken bevatten (IP-adressen worden zo uitgesloten).
    """
    import re as _re2
    if "#" in name:
        return name.split("#")[0].strip()
    m = _re2.match(r'^(.+?)\.(\d+)$', name)
    if m:
        prefix = m.group(1).strip()
        if _re2.search(r'[^0-9.]', prefix):
            return prefix
    return name


def _build_duplicate_ip_set(data: dict) -> set:
    """
    Set van IP-adressen die een echt conflict vormen.
    Uitgesloten: devices waarvan de naam een '#' bevat en het deel vóór '#'
    identiek is (= zelfde redundantiegroep / stack). Die mogen hetzelfde
    management-IP delen.
    Voorbeeld: SW01#1 en SW01#2 → groep 'SW01' → geen conflict.
    """
    from collections import defaultdict
    ip_to_groups: dict = defaultdict(set)
    for dv in data.get("devices", []):
        ip = _normalize_ip(dv.get("ip", ""))
        if not ip or ip == "—":
            continue
        name = dv.get("name", "")
        ip_to_groups[ip].add(_stack_group(name))
    # Conflict = IP gedeeld door minstens 2 verschillende groepen
    return {ip for ip, groups in ip_to_groups.items() if len(groups) > 1}


# ===========================================================================
# HEADER / FOOTER
# ===========================================================================

def _add_header(doc, org_name: str, version: str) -> None:
    hdr = doc.sections[0].header
    doc.sections[0].header_distance = Pt(28)
    for p in list(hdr.paragraphs):
        p._element.getparent().remove(p._element)

    p = hdr.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r1 = p.add_run("Netwerkrapport")
    r1.bold = True; r1.font.size = Pt(16); r1.font.color.rgb = _C_ZWART
    r2 = p.add_run(f"   |   {org_name}")
    r2.font.size = Pt(11); r2.font.color.rgb = _C_SUBTXT
    p.add_run("\t")
    rd = p.add_run(datetime.date.today().strftime("%d/%m/%Y"))
    rd.font.size = Pt(10); rd.font.color.rgb = _C_SUBTXT
    pPr  = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right"); tab.set(qn("w:pos"), "15840")
    tabs.append(tab); pPr.append(tabs)

    p_sep = hdr.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(2)
    p_sep.paragraph_format.space_after  = Pt(0)
    _add_border_bottom(p_sep, _C_ACCENT, 12)


def _add_footer(doc, version: str) -> None:
    ftr = doc.sections[0].footer
    for p in list(ftr.paragraphs):
        p._element.getparent().remove(p._element)

    p = ftr.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    rl = p.add_run(
        f"Networkmap Creator {version}  |  "
        f"Gegenereerd op {datetime.date.today().strftime('%d/%m/%Y')}"
    )
    rl.font.size = Pt(9); rl.font.color.rgb = _C_SUBTXT
    p.add_run("\t")

    def _field(ft: str) -> None:
        r = p.add_run()
        fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "begin"); r._r.append(fld)
        r2  = p.add_run()
        ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve"); ins.text = f" {ft} "; r2._r.append(ins)
        r3  = p.add_run()
        fl2 = OxmlElement("w:fldChar"); fl2.set(qn("w:fldCharType"), "separate"); r3._r.append(fl2)
        r4  = p.add_run("1"); r4.font.size = Pt(9); r4.font.color.rgb = _C_SUBTXT
        r5  = p.add_run()
        fl3 = OxmlElement("w:fldChar"); fl3.set(qn("w:fldCharType"), "end"); r5._r.append(fl3)

    rpre = p.add_run("Pagina "); rpre.font.size = Pt(9); rpre.font.color.rgb = _C_SUBTXT
    _field("PAGE")
    rmid = p.add_run(" van "); rmid.font.size = Pt(9); rmid.font.color.rgb = _C_SUBTXT
    _field("NUMPAGES")
    pPr  = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right"); tab.set(qn("w:pos"), "15840")
    tabs.append(tab); pPr.append(tabs)


# ===========================================================================
# SECTIE-TITELS (volledige breedte)
# ===========================================================================

def _add_site_header(doc, site_name: str, location: str = "") -> None:
    cell = _full_cell(doc, _rgb_hex(_C_ACCENT))
    _clear_p(cell)
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"📍  {site_name}")
    r1.bold = True; r1.font.size = Pt(16); r1.font.color.rgb = _C_WIT
    if location:
        r2 = p.add_run(f"   —   {location}")
        r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)
    _spacer(doc, 0.4)


def _add_room_header(doc, room_name: str, floor: str = "") -> None:
    cell = _full_cell(doc, _rgb_hex(_C_GRIJSDK), top=60, bottom=60)
    _clear_p(cell)
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"🚪  {room_name}")
    r1.bold = True; r1.font.size = Pt(13); r1.font.color.rgb = _C_ACCENT
    if floor:
        r2 = p.add_run(f"   ·   {floor}")
        r2.font.size = Pt(10); r2.font.color.rgb = _C_SUBTXT
    _spacer(doc, 0.3)


def _add_rack_badge(doc, rack_name: str, room_name: str) -> None:
    cell = _full_cell(doc, "E8F0FA", top=60, bottom=60)
    _clear_p(cell)
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("🗄  "); r1.font.size = Pt(11); r1.font.color.rgb = _C_ACCENT
    r2 = p.add_run(rack_name); r2.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = _C_ACCENT
    r3 = p.add_run(f"   ·   {room_name}"); r3.font.size = Pt(9); r3.font.color.rgb = _C_SUBTXT
    _spacer(doc, 0.2)


def _add_section_label(doc, label: str) -> None:
    cell = _full_cell(doc, _rgb_hex(_C_GRIJSDK), top=40, bottom=40)
    _clear_p(cell)
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label)
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _C_ACCENT


# ===========================================================================
# TITELBLAD
# ===========================================================================

def _build_titlepage(doc, data: dict, version: str) -> None:
    # Compacte cover: alles op één A4-landscape pagina (17.2cm beschikbaar)
    _spacer(doc, 0.4)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Networkmap Creator"); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = _C_ACCENT
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run("Netwerkinfrastructuur Rapport"); r2.font.size = Pt(13); r2.font.color.rgb = _C_SUBTXT
    pl = doc.add_paragraph(); pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pl.paragraph_format.space_before = Pt(2); pl.paragraph_format.space_after = Pt(2)
    _add_border_bottom(pl, _C_ACCENT, 12)
    pd = doc.add_paragraph(); pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pd.paragraph_format.space_before = Pt(0); pd.paragraph_format.space_after = Pt(0)
    r3 = pd.add_run(f"Gegenereerd op: {datetime.date.today().strftime('%d %B %Y')}")
    r3.font.size = Pt(10); r3.font.color.rgb = _C_SUBTXT
    pv = doc.add_paragraph(); pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pv.paragraph_format.space_before = Pt(0); pv.paragraph_format.space_after = Pt(2)
    r4 = pv.add_run(f"Versie software: {version}")
    r4.font.size = Pt(9); r4.italic = True; r4.font.color.rgb = _C_SUBTXT
    # Validatiestatus
    _val_status = data.get("_validation_status", "TE CONTROLEREN")
    _val_color  = {"GEVALIDEERD": _C_GROEN, "TE CONTROLEREN": _C_ORANJE,
                   "NIET DEFINITIEF": _C_ROOD, "ONGELDIG": _C_ROOD}.get(_val_status, _C_ORANJE)
    pst = doc.add_paragraph(); pst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pst.paragraph_format.space_before = Pt(0); pst.paragraph_format.space_after = Pt(4)
    rst = pst.add_run(f"Validatiestatus: {_val_status}")
    rst.bold = True; rst.font.size = Pt(10); rst.font.color.rgb = _val_color

    # F3 — bedrijfsblok (enkel bij scope van één bedrijf)
    _company = data.get("_company")
    if _company:
        _cname = _company.get("name", "")
        if _cname:
            pcn = doc.add_paragraph(); pcn.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pcn.paragraph_format.space_before = Pt(2); pcn.paragraph_format.space_after = Pt(0)
            rcn = pcn.add_run(_cname); rcn.bold = True
            rcn.font.size = Pt(14); rcn.font.color.rgb = _C_ACCENT
        _caddr = _company.get("address", "")
        if _caddr:
            pca = doc.add_paragraph(); pca.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pca.paragraph_format.space_before = Pt(0); pca.paragraph_format.space_after = Pt(0)
            rca = pca.add_run(_caddr); rca.font.size = Pt(9); rca.font.color.rgb = _C_SUBTXT
        _bits = []
        if _company.get("vat"):     _bits.append(f"BTW {_company['vat']}")
        if _company.get("phone"):   _bits.append(_company["phone"])
        if _company.get("email"):   _bits.append(_company["email"])
        if _company.get("website"): _bits.append(_company["website"])
        if _bits:
            pcc = doc.add_paragraph(); pcc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pcc.paragraph_format.space_before = Pt(0); pcc.paragraph_format.space_after = Pt(4)
            rcc = pcc.add_run("   •   ".join(_bits))
            rcc.font.size = Pt(9); rcc.font.color.rgb = _C_SUBTXT

    sites   = get_all_sites(data)
    n_vlans = data.get("_vlan_count_corrected", 0)
    _ap_counts = data.get("_action_counts", {})

    # Stats + VLAN + actiepunten in één gecombineerde tabel
    stats = [
        ("📍 Sites",        str(len(sites))),
        ("🚪 Ruimtes",      str(sum(len(s.get("rooms", [])) for s in sites))),
        ("🗄 Racks",        str(sum(len(r.get("racks", [])) for s in sites for r in s.get("rooms", [])))),
        ("💻 Devices",      str(len(data.get("devices", [])))),
        ("⬡ Poorten",           str(len(data.get("ports", [])))),
        ("🔗 Verbindingen",  str(len(data.get("connections", [])))),
        ("🌐 Wandpunten",    str(sum(len(r.get("wall_outlets", [])) for s in sites for r in s.get("rooms", [])))),
        ("🔷 VLANs gebruikt",        str(n_vlans)),
        ("🔷 VLANs gedefinieerd",    str(data.get("_vlan_count_defined", n_vlans))),
        ("🔷 VLANs volledig",        str(data.get("_vlan_count_complete", "—"))),
        ("⚠️  VLAN aandachtspunten",   str(data.get("_vlan_issues", "—"))),
    ]
    if _ap_counts:
        stats += [
            ("📋 Open actiepunten",
             f"{_ap_counts.get('total', 0)}  —  Hoog: {_ap_counts.get('high', 0)}  •  "
             f"Middel: {_ap_counts.get('medium', 0)}  •  Laag: {_ap_counts.get('low', 0)}"),
        ]
    tbl = doc.add_table(rows=0, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _col_widths(tbl, [Cm(8), Cm(11)])
    for i, (label, val) in enumerate(stats):
        row = tbl.add_row(); _cant_split(row)
        c0, c1 = row.cells[0], row.cells[1]
        bg = _rgb_hex(_C_GRIJS) if i % 2 == 0 else "FFFFFF"
        for cell in (c0, c1):
            _shade(cell, bg); _no_borders(cell)
        _margins(c0, top=30, bottom=30, left=150, right=40)
        _margins(c1, top=30, bottom=30, left=40, right=150)
        is_action = label.startswith("📋")
        val_color = _val_color if is_action and _ap_counts.get('high', 0) > 0 else _C_ACCENT
        _cell_p(c0, label, size=9, color=_C_ZWART, bold=is_action)
        _cell_p(c1, val, bold=is_action, size=9, color=val_color)


# ===========================================================================
# RACK BEZETTINGS-OVERZICHT
# ===========================================================================

def _add_rack_overview(doc, room: dict) -> None:
    racks = room.get("racks", [])
    if not racks:
        return
    _add_section_label(doc, "Rack bezetting")
    COLS = ["Rack", "Tot. U", "Bezet U", "Vrij U", "Bezetting"]
    W    = [Cm(5), Cm(2.5), Cm(2.5), Cm(2.5), Cm(15.2)]
    tbl  = doc.add_table(rows=1, cols=5); tbl.style = "Table Grid"
    hdr  = tbl.rows[0]; _tbl_header(hdr)
    for i, (col, w) in enumerate(zip(COLS, W)):
        cell = hdr.cells[i]; cell.width = w
        _shade(cell, _rgb_hex(_C_ACCENT)); _no_borders(cell); _margins(cell, top=50, bottom=50, left=120, right=80)
        _cell_p(cell, col, bold=True, size=9, color=_C_WIT)

    for ri, rack in enumerate(racks):
        total = rack.get("total_units", 42)
        used  = sum(max(1, s.get("height", 1)) for s in rack.get("slots", []))
        used  = min(used, total)
        free  = total - used
        pct   = int(used / total * 100) if total else 0
        bar   = "█" * (pct // 10) + "░" * (10 - pct // 10)
        bc    = (_C_ROOD if pct >= 80 else _C_ORANJE if pct >= 60 else _C_GROEN)
        row   = tbl.add_row(); _cant_split(row)
        bg    = _rgb_hex(_C_GRIJS) if ri % 2 == 0 else "FFFFFF"
        for i, (val, w, col) in enumerate(zip(
                [rack["name"], str(total), str(used), str(free), f"{pct}%  {bar}"], W,
                [_C_ZWART, _C_SUBTXT, _C_SUBTXT, _C_SUBTXT, bc])):
            cell = row.cells[i]; cell.width = w
            _shade(cell, bg); _no_borders(cell); _margins(cell, top=40, bottom=40, left=120, right=80)
            _cell_p(cell, val, bold=(i == 0), size=9, color=col)

    _col_widths(tbl, W); _spacer(doc, 0.3)
    n_wo = len(room.get("wall_outlets", []))
    if n_wo:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"  {n_wo} wandpunt(en) geconfigureerd.")
        r.font.size = Pt(9); r.font.color.rgb = _C_SUBTXT; r.italic = True


# ===========================================================================
# DEVICE TABEL
# ===========================================================================

def _add_device_table(doc, rack: dict, idx: dict) -> None:
    slots = sorted(rack.get("slots", []), key=lambda s: s.get("u_start", 0), reverse=True)
    if not slots:
        return
    COLS  = ["Device", "Type", "Merk / Model", "IP", "Slot", "Notitie"]
    W     = [Cm(4.5), Cm(3.0), Cm(4.5), Cm(3.5), Cm(1.5), Cm(10.7)]
    NCOLS = 6
    tbl   = doc.add_table(rows=0, cols=NCOLS); tbl.style = "Table Grid"

    # Rij 0: sectielabel (merged, tblHeader)
    lbl_row = tbl.add_row(); _tbl_header(lbl_row)
    lbl_row.cells[0].merge(lbl_row.cells[NCOLS - 1])
    cell = lbl_row.cells[0]
    _shade(cell, _rgb_hex(_C_GRIJSDK)); _no_borders(cell); _margins(cell, top=40, bottom=40, left=200, right=200)
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r = p.add_run("💻  Devices"); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _C_ACCENT

    # Rij 1: kolomheaders (tblHeader)
    hdr_row = tbl.add_row(); _tbl_header(hdr_row)
    for i, (col, w) in enumerate(zip(COLS, W)):
        cell = hdr_row.cells[i]; cell.width = w
        _shade(cell, _rgb_hex(_C_GRIJSDK)); _no_borders(cell); _margins(cell, top=50, bottom=50, left=120, right=80)
        _cell_p(cell, col, bold=True, size=9, color=_C_ACCENT)

    total_u = rack.get("total_units", 42)
    for ri, slot in enumerate(slots):
        dev = idx["dev"].get(slot.get("device_id", ""))
        if not dev:
            continue
        bm   = " ".join(filter(None, [dev.get("brand", ""), dev.get("model", "")])) or "—"
        ip   = _normalize_ip(dev.get("ip", "")) or "—"
        notes = dev.get("notes", "") or "—"
        u    = slot.get("u_start")
        slot_lbl = str(total_u - u + 1) if u is not None else "—"
        tl   = _TYPE_MAP.get(dev.get("type", ""), dev.get("type", "") or "—")
        try:
            from app.helpers.i18n import t as _t
            lbl = _t(f"device_{dev.get('type','')}")
            if lbl and not lbl.startswith("["):
                tl = lbl
        except Exception:
            pass
        row = tbl.add_row(); _cant_split(row)
        bg  = _rgb_hex(_C_GRIJS) if ri % 2 == 0 else "FFFFFF"
        for i, (val, w, col) in enumerate(zip(
                [dev["name"], tl, bm, ip, slot_lbl, notes], W,
                [_C_ZWART, _C_SUBTXT, _C_SUBTXT, _C_SUBTXT, _C_SUBTXT, _C_SUBTXT])):
            cell = row.cells[i]; cell.width = w
            _shade(cell, bg); _no_borders(cell); _margins(cell, top=40, bottom=40, left=120, right=80)
            _cell_p(cell, val, bold=(i == 0), size=9, color=col)
    _col_widths(tbl, W); _spacer(doc, 0.3)


# ===========================================================================
# POORTENTABEL (per device, device-label als merged tblHeader rij)
# ===========================================================================

def _add_port_table(doc, rack: dict, data: dict, idx: dict) -> None:
    slots = sorted(rack.get("slots", []), key=lambda s: s.get("u_start", 0), reverse=True)
    COLS  = ["Poort", "Zijde", "VLAN", "Verbonden met"]
    W     = [Cm(2.5), Cm(2.0), Cm(1.5), Cm(21.7)]
    NCOLS = 4
    first_device = True

    for slot in slots:
        dev = idx["dev"].get(slot.get("device_id", ""))
        if not dev:
            continue
        pf = sorted([p for p in data.get("ports", []) if p["device_id"] == dev["id"] and p["side"] == "front"], key=lambda p: p["number"])
        pb = sorted([p for p in data.get("ports", []) if p["device_id"] == dev["id"] and p["side"] == "back"],  key=lambda p: p["number"])
        all_ports = pf + pb
        if not all_ports:
            continue

        tbl = doc.add_table(rows=0, cols=NCOLS); tbl.style = "Table Grid"

        # Rij 0 (eerste device): "⬡ Poortoverzicht" sectielabel
        if first_device:
            sec_row = tbl.add_row(); _tbl_header(sec_row)
            sec_row.cells[0].merge(sec_row.cells[NCOLS - 1])
            cell = sec_row.cells[0]
            _shade(cell, _rgb_hex(_C_GRIJSDK)); _no_borders(cell); _margins(cell, top=40, bottom=40, left=200, right=200)
            _clear_p(cell); p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
            r = p.add_run("⬡  Poortoverzicht"); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _C_ACCENT
            first_device = False

        # Device-label rij (merged, tblHeader)
        lbl_row = tbl.add_row(); _tbl_header(lbl_row)
        lbl_row.cells[0].merge(lbl_row.cells[NCOLS - 1])
        cell = lbl_row.cells[0]
        _shade(cell, "F0F4FA"); _no_borders(cell); _margins(cell, top=35, bottom=35, left=200, right=200)
        _clear_p(cell); p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"  💻  {dev['name']}"); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _C_ACCENT
        dev_loc = idx.get("loc", {}).get(dev["id"])
        if dev_loc:
            r2 = p.add_run(f"   ·   {dev_loc['room']}  /  {dev_loc['rack']}  /  U{dev_loc['slot']}")
            r2.font.size = Pt(8); r2.font.color.rgb = _C_SUBTXT

        # Kolomheaders (tblHeader)
        hdr_row = tbl.add_row(); _tbl_header(hdr_row)
        for i, (col, w) in enumerate(zip(COLS, W)):
            cell = hdr_row.cells[i]; cell.width = w
            _shade(cell, _rgb_hex(_C_GRIJSDK)); _no_borders(cell); _margins(cell, top=40, bottom=40, left=120, right=80)
            _cell_p(cell, col, bold=True, size=8, color=_C_ACCENT)

        # Data rijen
        for pi, port in enumerate(all_ports):
            side_str = "VOOR" if port["side"] == "front" else "ACHTER"
            vlan_str = str(port.get("vlan") or "—")
            dest     = _conn_label(data, idx, port["id"])
            row      = tbl.add_row(); _cant_split(row)
            bg       = _rgb_hex(_C_GRIJS) if pi % 2 == 0 else "FFFFFF"
            dest_color = (
                _C_GROEN if dest.startswith("🌐") else
                _C_BLAUW if dest.startswith("🖥") else
                _C_ACCENT if dest != "—" else _C_SUBTXT
            )
            for i, (val, w, col) in enumerate(zip(
                    [port.get("name", f"Port {port['number']}"), side_str, vlan_str, dest], W,
                    [_C_ZWART, _C_SUBTXT, _C_SUBTXT, dest_color])):
                cell = row.cells[i]; cell.width = w
                _shade(cell, bg); _no_borders(cell); _margins(cell, top=30, bottom=30, left=120, right=80)
                _cell_p(cell, val, size=8, color=col)
        _col_widths(tbl, W); _spacer(doc, 0.25)
    _spacer(doc, 0.3)


# ===========================================================================
# PER-SITE SECTIES: WANDPUNTEN
# ===========================================================================

def _add_site_outlets(doc, data: dict, idx: dict, site: dict) -> None:
    sn = site["name"]
    rooms_with_outlets = [r for r in site.get("rooms", []) if r.get("wall_outlets")]
    if not rooms_with_outlets:
        return
    _add_page_break(doc)
    cell = _full_cell(doc, _rgb_hex(_C_ACCENT))
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"🌐  Wandpunten  ·  {sn}")
    r1.bold = True; r1.font.size = Pt(14); r1.font.color.rgb = _C_WIT
    _spacer(doc, 0.4)

    COLS  = ["Naam", "VLAN", "Eindapparaat", "Verbonden met (poort)"]
    W     = [Cm(2.5), Cm(1.5), Cm(5.0), Cm(18.7)]
    NCOLS = 4

    for room in rooms_with_outlets:
        outlets = room.get("wall_outlets", [])
        cell = _full_cell(doc, _rgb_hex(_C_GRIJSDK), top=60, bottom=60)
        _clear_p(cell); p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run(f"🚪  {room['name']}")
        r1.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = _C_ACCENT
        r2 = p.add_run(f"   ·   {len(outlets)} wandpunten")
        r2.font.size = Pt(9); r2.font.color.rgb = _C_SUBTXT
        if room.get("floor"):
            r3 = p.add_run(f"   ·   {room['floor']}"); r3.font.size = Pt(9); r3.font.color.rgb = _C_SUBTXT
        _spacer(doc, 0.2)

        from collections import defaultdict
        loc_groups: dict = defaultdict(list)
        for wo in outlets:
            loc_groups[wo.get("location_description", "") or ""].append(wo)
        sorted_groups = sorted(loc_groups.items(), key=lambda kv: _loc_group_label(kv[0]))

        tbl = doc.add_table(rows=0, cols=NCOLS); tbl.style = "Table Grid"
        sec_row = tbl.add_row(); _tbl_header(sec_row)
        sec_row.cells[0].merge(sec_row.cells[NCOLS - 1]); cell = sec_row.cells[0]
        _shade(cell, _rgb_hex(_C_GRIJSDK)); _no_borders(cell); _margins(cell, top=40, bottom=40, left=200, right=200)
        _clear_p(cell); p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"🚪  {room['name']}  ·  {sn}"); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _C_ACCENT

        hdr_row = tbl.add_row(); _tbl_header(hdr_row)
        for i, (col, w) in enumerate(zip(COLS, W)):
            cell = hdr_row.cells[i]; cell.width = w
            _shade(cell, _rgb_hex(_C_ACCENT)); _no_borders(cell); _margins(cell, top=50, bottom=50, left=120, right=80)
            _cell_p(cell, col, bold=True, size=9, color=_C_WIT)

        row_idx = 0
        for loc_key, wo_list in sorted_groups:
            loc_label = _loc_group_label(loc_key)
            sub_row   = tbl.add_row(); _cant_split(sub_row)
            sub_row.cells[0].merge(sub_row.cells[NCOLS - 1]); cell = sub_row.cells[0]
            _shade(cell, "EEF3FA"); _no_borders(cell); _margins(cell, top=30, bottom=30, left=300, right=200)
            _clear_p(cell); p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
            r = p.add_run(f"  📌  {loc_label}"); r.bold = True; r.font.size = Pt(8); r.font.color.rgb = _C_ACCENT

            for wo in sorted(wo_list, key=lambda w: w.get("name", "")):
                ep      = idx["ep"].get(wo.get("endpoint_id", ""))
                ep_name = ep["name"] if ep else "—"
                vlan    = str(wo.get("vlan") or "—")
                dest    = "—"
                for conn in data.get("connections", []):
                    if conn.get("from_id") == wo["id"] or conn.get("to_id") == wo["id"]:
                        oid = conn["to_id"] if conn["from_id"] == wo["id"] else conn["from_id"]
                        ot  = conn["to_type"] if conn["from_id"] == wo["id"] else conn["from_type"]
                        if ot in ("port", None):
                            p2  = idx["port"].get(oid)
                            dev = idx["dev"].get(p2["device_id"]) if p2 else None
                            if p2 and dev:
                                side = "F" if p2["side"] == "front" else "B"
                                dest = f"{dev['name']} / {p2['name']} ({side}){_dev_loc_str(idx, dev['id'])}"
                        break
                row = tbl.add_row(); _cant_split(row)
                bg  = _rgb_hex(_C_GRIJS) if row_idx % 2 == 0 else "FFFFFF"; row_idx += 1
                ep_color = _C_BLAUW if ep else _C_SUBTXT
                for i, (val, w, col) in enumerate(zip(
                        [wo.get("name", wo["id"]), vlan, ep_name, dest], W,
                        [_C_ZWART, _C_SUBTXT, ep_color, _C_ACCENT if dest != "—" else _C_SUBTXT])):
                    cell = row.cells[i]; cell.width = w
                    _shade(cell, bg); _no_borders(cell); _margins(cell, top=35, bottom=35, left=120, right=80)
                    _cell_p(cell, val, bold=(i == 0), size=9, color=col)

        _col_widths(tbl, W); _spacer(doc, 0.4)


# ===========================================================================
# PER-SITE SECTIES: VERBINDINGEN
# ===========================================================================

def _add_site_connections(doc, data: dict, idx: dict, site: dict) -> None:
    site_dev_ids = {
        slot.get("device_id")
        for room in site.get("rooms", [])
        for rack in room.get("racks", [])
        for slot in rack.get("slots", [])
        if slot.get("device_id")
    }
    site_wo_ids = {wo["id"] for room in site.get("rooms", []) for wo in room.get("wall_outlets", [])}
    conns = [c for c in data.get("connections", [])
             if (_is_site_port(c.get("from_id"), c.get("from_type"), idx, site_dev_ids) or
                 c.get("from_id") in site_wo_ids or
                 _is_site_port(c.get("to_id"), c.get("to_type"), idx, site_dev_ids) or
                 c.get("to_id") in site_wo_ids)]
    if not conns:
        return

    _add_page_break(doc)
    cell = _full_cell(doc, _rgb_hex(_C_ACCENT))
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"🔗  Verbindingen  ·  {site['name']}")
    r1.bold = True; r1.font.size = Pt(14); r1.font.color.rgb = _C_WIT
    r2 = p.add_run(f"   —   {len(conns)} totaal")
    r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)
    _spacer(doc, 0.4)

    COLS  = ["Van", "Naar", "Kabeltype", "VLAN", "Notitie"]
    W     = [Cm(8.5), Cm(8.5), Cm(2.5), Cm(1.5), Cm(6.7)]
    NCOLS = 5
    tbl   = doc.add_table(rows=0, cols=NCOLS); tbl.style = "Table Grid"

    sec_row = tbl.add_row(); _tbl_header(sec_row)
    sec_row.cells[0].merge(sec_row.cells[NCOLS - 1]); cell = sec_row.cells[0]
    _shade(cell, _rgb_hex(_C_GRIJSDK)); _no_borders(cell); _margins(cell, top=40, bottom=40, left=200, right=200)
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"🔗  Verbindingen  ·  {site['name']}"); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _C_ACCENT

    hdr_row = tbl.add_row(); _tbl_header(hdr_row)
    for i, (col, w) in enumerate(zip(COLS, W)):
        cell = hdr_row.cells[i]; cell.width = w
        _shade(cell, _rgb_hex(_C_ACCENT)); _no_borders(cell); _margins(cell, top=50, bottom=50, left=120, right=80)
        _cell_p(cell, col, bold=True, size=9, color=_C_WIT)

    def _dc(val: str) -> RGBColor:
        if val.startswith("🌐"): return _C_GROEN
        if val.startswith("🖥"): return _C_BLAUW
        return _C_ACCENT if val != "—" else _C_SUBTXT

    for ci, conn in enumerate(conns):
        fl    = _resolve_side(data, idx, conn.get("from_id"), conn.get("from_type", "port"))
        tl    = _resolve_side(data, idx, conn.get("to_id"),   conn.get("to_type",   "port"))
        cable = _CABLE_LABELS.get(conn.get("cable_type", ""), conn.get("cable_type") or "—")
        note  = conn.get("notes") or "—"
        vlan_str = "—"
        for eid, etype in [(conn.get("from_id"), conn.get("from_type", "port")),
                           (conn.get("to_id"),   conn.get("to_type",   "port"))]:
            if etype in ("port", None):
                pt = idx["port"].get(eid)
                if pt and pt.get("vlan"):
                    vlan_str = str(pt["vlan"]); break

        row = tbl.add_row(); _cant_split(row)
        bg  = _rgb_hex(_C_GRIJS) if ci % 2 == 0 else "FFFFFF"
        for i, (val, w, col) in enumerate(zip(
                [fl, tl, cable, vlan_str, note], W,
                [_dc(fl), _dc(tl), _C_SUBTXT, _C_SUBTXT, _C_SUBTXT])):
            cell = row.cells[i]; cell.width = w
            _shade(cell, bg); _no_borders(cell); _margins(cell, top=35, bottom=35, left=120, right=80)
            _cell_p(cell, val, size=9, color=col)

    _col_widths(tbl, W)


# ===========================================================================
# PER-SITE SECTIES: DIRECT VERBONDEN APPARATEN
# ===========================================================================

def _add_site_direct_endpoints(doc, data: dict, idx: dict, site: dict) -> None:
    site_dev_ids = {
        slot.get("device_id")
        for room in site.get("rooms", [])
        for rack in room.get("racks", [])
        for slot in rack.get("slots", [])
        if slot.get("device_id")
    }
    direct_conns = [c for c in data.get("connections", [])
                    if c.get("to_type") == "endpoint" or c.get("from_type") == "endpoint"]
    rows_data = []
    for conn in direct_conns:
        port_id = conn["from_id"] if conn.get("to_type") == "endpoint" else conn["to_id"]
        ep_id   = conn["to_id"]   if conn.get("to_type") == "endpoint" else conn["from_id"]
        port    = idx["port"].get(port_id)
        if not port or port.get("device_id") not in site_dev_ids:
            continue
        ep  = idx["ep"].get(ep_id)
        dev = idx["dev"].get(port.get("device_id", ""))
        if ep:
            rows_data.append((ep, port, dev, conn))
    if not rows_data:
        return

    _add_page_break(doc)
    cell = _full_cell(doc, _rgb_hex(_C_ACCENT))
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"🖥  Direct verbonden  ·  {site['name']}")
    r1.bold = True; r1.font.size = Pt(14); r1.font.color.rgb = _C_WIT
    r2 = p.add_run(f"   —   {len(rows_data)} apparaten")
    r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)
    _spacer(doc, 0.4)

    COLS  = ["Naam", "Type", "Locatie", "Kabeltype", "Verbonden poort"]
    W     = [Cm(5.0), Cm(3.5), Cm(5.0), Cm(2.5), Cm(11.7)]
    NCOLS = 5
    tbl   = doc.add_table(rows=0, cols=NCOLS); tbl.style = "Table Grid"

    sec_row = tbl.add_row(); _tbl_header(sec_row)
    sec_row.cells[0].merge(sec_row.cells[NCOLS - 1]); cell = sec_row.cells[0]
    _shade(cell, _rgb_hex(_C_GRIJSDK)); _no_borders(cell); _margins(cell, top=40, bottom=40, left=200, right=200)
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"🖥  Direct verbonden  ·  {site['name']}"); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _C_ACCENT

    hdr_row = tbl.add_row(); _tbl_header(hdr_row)
    for i, (col, w) in enumerate(zip(COLS, W)):
        cell = hdr_row.cells[i]; cell.width = w
        _shade(cell, _rgb_hex(_C_ACCENT)); _no_borders(cell); _margins(cell, top=50, bottom=50, left=120, right=80)
        _cell_p(cell, col, bold=True, size=9, color=_C_WIT)

    for ri, (ep, port, dev, conn) in enumerate(sorted(rows_data, key=lambda r: r[0].get("name", ""))):
        side     = "F" if port["side"] == "front" else "B"
        loc_str  = _dev_loc_str(idx, dev["id"]) if dev else ""
        port_lbl = f"{dev['name']} / {port['name']} ({side}){loc_str}" if dev else port.get("name", "?")
        cable    = _CABLE_LABELS.get(conn.get("cable_type", ""), conn.get("cable_type") or "—")
        ep_type  = _EP_TYPE_MAP.get(ep.get("type", ""), ep.get("type", "") or "—")
        row = tbl.add_row(); _cant_split(row)
        bg  = _rgb_hex(_C_GRIJS) if ri % 2 == 0 else "FFFFFF"
        for i, (val, w, col) in enumerate(zip(
                [ep.get("name", "?"), ep_type, ep.get("location", "") or "—", cable, port_lbl], W,
                [_C_ZWART, _C_SUBTXT, _C_SUBTXT, _C_SUBTXT, _C_ACCENT])):
            cell = row.cells[i]; cell.width = w
            _shade(cell, bg); _no_borders(cell); _margins(cell, top=35, bottom=35, left=120, right=80)
            _cell_p(cell, val, bold=(i == 0), size=9, color=col)
    _col_widths(tbl, W)


# ===========================================================================
# PER-SITE SECTIES: VLAN
# ===========================================================================

def _add_site_vlans(doc, data: dict, idx: dict, site: dict) -> None:
    site_dev_ids = {
        slot.get("device_id")
        for room in site.get("rooms", [])
        for rack in room.get("racks", [])
        for slot in rack.get("slots", [])
        if slot.get("device_id")
    }
    vlan_map: dict = {}
    for p in data.get("ports", []):
        if p.get("device_id") not in site_dev_ids:
            continue
        v = p.get("vlan")
        if v:
            try: vi = int(v)
            except Exception: vi = v
            vlan_map.setdefault(vi, {"ports": [], "wall_outlets": []})["ports"].append(p)
    for room in site.get("rooms", []):
        for wo in room.get("wall_outlets", []):
            v = wo.get("vlan")
            if v:
                try: vi = int(v)
                except Exception: vi = v
                vlan_map.setdefault(vi, {"ports": [], "wall_outlets": []})["wall_outlets"].append((wo, room["name"]))
    if not vlan_map:
        return

    _add_page_break(doc)
    cell = _full_cell(doc, _rgb_hex(_C_ACCENT))
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"🔷  VLAN Overzicht  ·  {site['name']}")
    r1.bold = True; r1.font.size = Pt(14); r1.font.color.rgb = _C_WIT
    r2 = p.add_run(f"   —   {len(vlan_map)} VLAN{'s' if len(vlan_map) > 1 else ''}")
    r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)
    _spacer(doc, 0.4)

    for vlan_id in sorted(vlan_map.keys()):
        vdata = vlan_map[vlan_id]; ports = vdata["ports"]; wos = vdata["wall_outlets"]
        vcfg    = data.get("_vlan_config", {}).get(vlan_id, {})
        vname   = vcfg.get("name", "")
        vip     = vcfg.get("ip", "")
        vsubnet = vcfg.get("subnet", "")
        vdesc   = vcfg.get("description", "")

        cell = _full_cell(doc, "E8F0FA", top=60, bottom=60)
        _clear_p(cell); p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        vlan_title = f"🔷  VLAN {vlan_id}"
        if vname: vlan_title += f"  —  {vname}"
        r1 = p.add_run(vlan_title); r1.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = _C_ACCENT
        r2 = p.add_run(f"   ·   {len(ports)+len(wos)} items  ({len(ports)} poorten  /  {len(wos)} wandpunten)")
        r2.font.size = Pt(9); r2.font.color.rgb = _C_SUBTXT
        if vip or vsubnet or vdesc:
            p2 = cell.add_paragraph()
            p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(0)
            extra_parts = []
            if vip: extra_parts.append(f"Gateway: {vip}")
            if vsubnet: extra_parts.append(f"Subnet: {vsubnet}")
            if vdesc: extra_parts.append(vdesc)
            r3 = p2.add_run("  " + "   ·   ".join(extra_parts))
            r3.font.size = Pt(8); r3.font.color.rgb = _C_SUBTXT; r3.italic = True
        _spacer(doc, 0.2)

        if ports:
            COLS_P = ["Device", "Poort", "Zijde", "Locatie"]; W_P = [Cm(5.0), Cm(3.0), Cm(2.0), Cm(17.7)]
            tbl_p = doc.add_table(rows=0, cols=4); tbl_p.style = "Table Grid"
            sec_row = tbl_p.add_row(); _tbl_header(sec_row)
            sec_row.cells[0].merge(sec_row.cells[3]); cell = sec_row.cells[0]
            _shade(cell, _rgb_hex(_C_GRIJSDK)); _no_borders(cell); _margins(cell, top=35, bottom=35, left=200, right=200)
            _clear_p(cell); p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
            r = p.add_run("Poorten"); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _C_ACCENT
            hdr_row = tbl_p.add_row(); _tbl_header(hdr_row)
            for i, (col, w) in enumerate(zip(COLS_P, W_P)):
                cell = hdr_row.cells[i]; cell.width = w
                _shade(cell, _rgb_hex(_C_GRIJSDK)); _no_borders(cell); _margins(cell, top=40, bottom=40, left=120, right=80)
                _cell_p(cell, col, bold=True, size=9, color=_C_ACCENT)
            for ri, port in enumerate(sorted(ports, key=lambda p: (idx["dev"].get(p.get("device_id", ""), {}).get("name", ""), p["number"]))):
                dev = idx["dev"].get(port.get("device_id", ""))
                side = "VOOR" if port["side"] == "front" else "ACHTER"
                loc_str = _dev_loc_str(idx, dev["id"]).strip(" ()") if dev else ""
                row = tbl_p.add_row(); _cant_split(row)
                bg  = _rgb_hex(_C_GRIJS) if ri % 2 == 0 else "FFFFFF"
                for i, (val, w, col) in enumerate(zip(
                        [dev["name"] if dev else "—", port.get("name", "?"), side, loc_str], W_P,
                        [_C_ZWART, _C_SUBTXT, _C_SUBTXT, _C_SUBTXT])):
                    cell = row.cells[i]; cell.width = w
                    _shade(cell, bg); _no_borders(cell); _margins(cell, top=35, bottom=35, left=120, right=80)
                    _cell_p(cell, val, bold=(i == 0), size=9, color=col)
            _col_widths(tbl_p, W_P); _spacer(doc, 0.2)

        if wos:
            COLS_W = ["Wandpunt", "Locatie", "Ruimte", "Eindapparaat"]; W_W = [Cm(3.0), Cm(6.0), Cm(5.0), Cm(13.7)]
            tbl_w = doc.add_table(rows=0, cols=4); tbl_w.style = "Table Grid"
            sec_row = tbl_w.add_row(); _tbl_header(sec_row)
            sec_row.cells[0].merge(sec_row.cells[3]); cell = sec_row.cells[0]
            _shade(cell, _rgb_hex(_C_GRIJSDK)); _no_borders(cell); _margins(cell, top=35, bottom=35, left=200, right=200)
            _clear_p(cell); p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
            r = p.add_run("Wandpunten"); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _C_ACCENT
            hdr_row = tbl_w.add_row(); _tbl_header(hdr_row)
            for i, (col, w) in enumerate(zip(COLS_W, W_W)):
                cell = hdr_row.cells[i]; cell.width = w
                _shade(cell, _rgb_hex(_C_GRIJSDK)); _no_borders(cell); _margins(cell, top=40, bottom=40, left=120, right=80)
                _cell_p(cell, col, bold=True, size=9, color=_C_ACCENT)
            for wi, (wo, rn) in enumerate(sorted(wos, key=lambda x: x[0].get("name", ""))):
                ep = idx["ep"].get(wo.get("endpoint_id", "")); ep_name = ep["name"] if ep else "—"
                loc = _wo_location(wo); row = tbl_w.add_row(); _cant_split(row)
                bg = _rgb_hex(_C_GRIJS) if wi % 2 == 0 else "FFFFFF"; ep_col = _C_BLAUW if ep else _C_SUBTXT
                for i, (val, w, col) in enumerate(zip(
                        [wo.get("name", "?"), loc, rn, ep_name], W_W,
                        [_C_ZWART, _C_SUBTXT, _C_SUBTXT, ep_col])):
                    cell = row.cells[i]; cell.width = w
                    _shade(cell, bg); _no_borders(cell); _margins(cell, top=35, bottom=35, left=120, right=80)
                    _cell_p(cell, val, bold=(i == 0), size=9, color=col)
            _col_widths(tbl_w, W_W); _spacer(doc, 0.3)
        _spacer(doc, 0.3)


# ===========================================================================
# GLOBALE SECTIES
# ===========================================================================

def _add_site_summary(doc, data: dict, idx: dict, risk_by_site: dict) -> None:
    """Site-samenvatting tabel."""
    _add_page_break(doc)
    cell = _full_cell(doc, _rgb_hex(_C_ACCENT))
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r = p.add_run("📊  Site-samenvatting"); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = _C_WIT
    _spacer(doc, 0.4)

    COLS = ["Site", "Locatie", "Racks", "Switches", "Patchpanels", "Wandpunten", "Verbindingen", "Risico WP"]
    W    = [Cm(4.0), Cm(5.5), Cm(1.8), Cm(2.2), Cm(2.5), Cm(2.5), Cm(2.8), Cm(2.4)]
    NCOLS = len(COLS)
    tbl   = doc.add_table(rows=0, cols=NCOLS); tbl.style = "Table Grid"
    hdr_row = tbl.add_row(); _tbl_header(hdr_row)
    for i, (col, w) in enumerate(zip(COLS, W)):
        cell = hdr_row.cells[i]; cell.width = w
        _shade(cell, _rgb_hex(_C_ACCENT)); _no_borders(cell); _margins(cell, top=50, bottom=50, left=80, right=60)
        _cell_p(cell, col, bold=True, size=9, color=_C_WIT)

    for ri, site in enumerate(get_all_sites(data)):
        sn = site["name"]
        site_dev_ids = {
            slot.get("device_id")
            for room in site.get("rooms", [])
            for rack in room.get("racks", [])
            for slot in rack.get("slots", [])
            if slot.get("device_id")
        }
        site_wo_ids = {wo["id"] for room in site.get("rooms", []) for wo in room.get("wall_outlets", [])}
        racks    = sum(1 for room in site.get("rooms", []) for _ in room.get("racks", []))
        switches = sum(1 for did in site_dev_ids if idx["dev"].get(did, {}).get("type") == "switch")
        patches  = sum(1 for did in site_dev_ids if idx["dev"].get(did, {}).get("type") in ("patch_panel", "patchpanel"))
        wo_count = sum(len(room.get("wall_outlets", [])) for room in site.get("rooms", []))
        conns    = sum(1 for c in data.get("connections", [])
                      if _is_site_port(c.get("from_id"), c.get("from_type"), idx, site_dev_ids) or
                         c.get("from_id") in site_wo_ids or
                         _is_site_port(c.get("to_id"), c.get("to_type"), idx, site_dev_ids) or
                         c.get("to_id") in site_wo_ids)
        risk     = len(risk_by_site.get(sn, []))
        row      = tbl.add_row(); _cant_split(row)
        bg       = _rgb_hex(_C_GRIJS) if ri % 2 == 0 else "FFFFFF"
        risk_col = _C_ROOD if risk > 0 else _C_GROEN
        for i, (val, w, col) in enumerate(zip(
                [sn, site.get("location", "—"), str(racks), str(switches),
                 str(patches), str(wo_count), str(conns), str(risk)], W,
                [_C_ZWART, _C_SUBTXT, _C_SUBTXT, _C_ACCENT, _C_SUBTXT, _C_SUBTXT, _C_SUBTXT, risk_col])):
            cell = row.cells[i]; cell.width = w
            _shade(cell, bg); _no_borders(cell); _margins(cell, top=40, bottom=40, left=80, right=60)
            _cell_p(cell, val, bold=(i == 0), size=9, color=col)
    _col_widths(tbl, W); _spacer(doc, 0.4)


def _add_glossary(doc) -> None:
    """Begrippenlijst."""
    cell = _full_cell(doc, _rgb_hex(_C_ACCENT))
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r = p.add_run("📖  Begrippen en legende"); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = _C_WIT
    _spacer(doc, 0.4)

    begrippen = [
        ("Device",           "Elk fysiek object in een rack: switch, router, firewall, patchpanel, kabelgoot, UPS."),
        ("Actief device",    "Netwerkapparaat met actieve elektronica: switch, router, firewall, access point, server."),
        ("Passief device",   "Niet-actief onderdeel: patchpanel of kabelgoot."),
        ("Eindapparaat",     "Toestel aan het einde van een verbinding: pc, laptop, printer, WAP, camera, machine, telefoon."),
        ("Wandpunt",         "Fysiek aansluitpunt in een ruimte, verbonden via kabel naar patchpanel of switch."),
        ("Poort",            "Fysieke aansluiting op een device. VOOR (F) = voorzijde rack. ACHTER (B) = achterzijde rack."),
        ("Verbinding",       "Geregistreerde kabelkoppeling tussen twee poorten, of tussen poort en wandpunt/eindapparaat."),
        ("Risico-WP",        "Wandpunt verbonden aan het netwerk zonder geregistreerd eindapparaat — onbeheerde toegang."),
        ("Onverbonden WP",   "Wandpunt zonder geregistreerde kabelverbinding. Fysieke controle aanbevolen."),
        ("Rack bezetting",   "Verhouding bezette U-posities t.o.v. totale rack-hoogte (1U ≈ 44,5mm)."),
        ("VLAN",             "Virtual LAN — logische netwerksegmentatie. Inter-VLAN communicatie vereist routing."),
        ("Uplink",           "Verbinding tussen twee actieve netwerkapparaten (switch→switch of switch→router)."),
        ("F / B",            "F = Front (voorzijde), B = Back (achterzijde)."),
        ("—",                "Niet ingevuld of niet van toepassing."),
    ]
    COLS = ["Begrip", "Betekenis"]; W = [Cm(4.5), Cm(23.2)]
    tbl  = doc.add_table(rows=0, cols=2); tbl.style = "Table Grid"
    hdr_row = tbl.add_row(); _tbl_header(hdr_row)
    for i, (col, w) in enumerate(zip(COLS, W)):
        cell = hdr_row.cells[i]; cell.width = w
        _shade(cell, _rgb_hex(_C_ACCENT)); _no_borders(cell); _margins(cell, top=50, bottom=50, left=120, right=80)
        _cell_p(cell, col, bold=True, size=9, color=_C_WIT)
    for ri, (term, uitleg) in enumerate(begrippen):
        row = tbl.add_row(); _cant_split(row)
        bg  = _rgb_hex(_C_GRIJS) if ri % 2 == 0 else "FFFFFF"
        for i, (val, w, col) in enumerate(zip([term, uitleg], W, [_C_ZWART, _C_SUBTXT])):
            cell = row.cells[i]; cell.width = w
            _shade(cell, bg); _no_borders(cell); _margins(cell, top=40, bottom=40, left=120, right=80)
            _cell_p(cell, val, bold=(i == 0), size=9, color=col)
    _col_widths(tbl, W); _spacer(doc, 0.4)


def _add_switch_overview(doc, data: dict, idx: dict, dup_ips: set) -> None:
    """Switch-overzicht met uplinks en IP-markering."""
    switches = [dv for dv in data.get("devices", []) if dv.get("type") == "switch"]
    if not switches:
        return
    _add_page_break(doc)
    cell = _full_cell(doc, _rgb_hex(_C_ACCENT))
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("🔀  Switch-overzicht"); r1.bold = True; r1.font.size = Pt(16); r1.font.color.rgb = _C_WIT
    r2 = p.add_run(f"   —   {len(switches)} switches")
    r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)
    _spacer(doc, 0.4)

    switch_ids = {dv["id"] for dv in switches}
    uplink_map: dict = {}
    for conn in data.get("connections", []):
        fp = idx["port"].get(conn.get("from_id"))
        tp = idx["port"].get(conn.get("to_id"))
        if not fp or not tp:
            continue
        fd = idx["dev"].get(fp.get("device_id", ""))
        td = idx["dev"].get(tp.get("device_id", ""))
        if fd and td and fd["id"] in switch_ids and td["id"] in switch_ids and fd["id"] != td["id"]:
            uplink_map.setdefault(fd["id"], []).append(f"{fp['name']}→{td['name']}/{tp['name']}")
            uplink_map.setdefault(td["id"], []).append(f"{tp['name']}→{fd['name']}/{fp['name']}")

    COLS  = ["Switch", "Site", "Ruimte / Rack / U", "Model", "IP", "MAC", "Uplinks"]
    W     = [Cm(3.5), Cm(3.5), Cm(5.0), Cm(3.5), Cm(3.0), Cm(4.0), Cm(5.2)]
    NCOLS = len(COLS)
    tbl   = doc.add_table(rows=0, cols=NCOLS); tbl.style = "Table Grid"

    sec_row = tbl.add_row(); _tbl_header(sec_row)
    sec_row.cells[0].merge(sec_row.cells[NCOLS - 1]); cell = sec_row.cells[0]
    _shade(cell, _rgb_hex(_C_GRIJSDK)); _no_borders(cell); _margins(cell, top=40, bottom=40, left=200, right=200)
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r = p.add_run("🔀  Alle switches"); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _C_ACCENT

    hdr_row = tbl.add_row(); _tbl_header(hdr_row)
    for i, (col, w) in enumerate(zip(COLS, W)):
        cell = hdr_row.cells[i]; cell.width = w
        _shade(cell, _rgb_hex(_C_ACCENT)); _no_borders(cell); _margins(cell, top=50, bottom=50, left=80, right=60)
        _cell_p(cell, col, bold=True, size=9, color=_C_WIT)

    for ri, dv in enumerate(sorted(switches, key=lambda d: (
            idx.get("loc", {}).get(d["id"], {}).get("site", ""), d["name"]))):
        loc     = idx.get("loc", {}).get(dv["id"])
        site_nm = loc["site"] if loc else "—"
        loc_str = f"{loc['room']} / {loc['rack']} / U{loc['slot']}" if loc else "—"
        bm      = " ".join(filter(None, [dv.get("brand", ""), dv.get("model", "")])) or "—"
        ip      = _normalize_ip(dv.get("ip", ""))
        mac     = _normalize_mac(dv.get("mac", ""))
        uplinks = ", ".join(uplink_map.get(dv["id"], [])[:3]) or "—"
        ip_dup  = ip in dup_ips and ip != "—"
        row     = tbl.add_row(); _cant_split(row)
        bg      = _rgb_hex(_C_GRIJS) if ri % 2 == 0 else "FFFFFF"
        ip_col  = _C_ORANJE if ip_dup else (_C_ACCENT if ip != "—" else _C_SUBTXT)
        for i, (val, w, col) in enumerate(zip(
                [dv["name"], site_nm, loc_str, bm, ip + (" ⚠" if ip_dup else ""), mac, uplinks], W,
                [_C_ZWART, _C_SUBTXT, _C_SUBTXT, _C_SUBTXT, ip_col, _C_SUBTXT,
                 _C_ACCENT if uplinks != "—" else _C_SUBTXT])):
            cell = row.cells[i]; cell.width = w
            _shade(cell, bg); _no_borders(cell); _margins(cell, top=35, bottom=35, left=80, right=60)
            _cell_p(cell, val, bold=(i == 0), size=9, color=col)
    _col_widths(tbl, W)
    if dup_ips:
        note_row = tbl.add_row()
        merged   = note_row.cells[0]
        for ci in range(1, len(COLS)):
            merged = merged.merge(note_row.cells[ci])
        _shade(merged, "FFF3CD"); _no_borders(merged)
        _margins(merged, top=40, bottom=40, left=120, right=80)
        _cell_p(merged,
                f"  ⚠  Dubbele IP-adressen (conflict): {', '.join(sorted(dup_ips))}",
                size=9, color=_C_ORANJE, bold=False)
    # Geen trailing spacer hier: _add_uplink_overview begint met _add_page_break;
    # een spacer net op paginagrens veroorzaakte een lege pagina 6.


def _add_uplink_overview(doc, data: dict, idx: dict) -> None:
    """Switch-naar-switch uplinks."""
    switch_ids = {dv["id"] for dv in data.get("devices", []) if dv.get("type") == "switch"}
    uplinks = []
    seen: set = set()
    for conn in data.get("connections", []):
        fp = idx["port"].get(conn.get("from_id"))
        tp = idx["port"].get(conn.get("to_id"))
        if not fp or not tp:
            continue
        fd = idx["dev"].get(fp.get("device_id", ""))
        td = idx["dev"].get(tp.get("device_id", ""))
        if fd and td and fd["id"] in switch_ids and td["id"] in switch_ids and fd["id"] != td["id"]:
            key = tuple(sorted([conn.get("from_id", ""), conn.get("to_id", "")]))
            if key not in seen:
                seen.add(key)
                cable = _CABLE_LABELS.get(conn.get("cable_type", ""), conn.get("cable_type") or "—")
                _sfp = "SFP" in (fp.get("name", "") + tp.get("name", "")).upper()
                _utp = any(kw in cable.upper() for kw in ("UTP", "CAT"))
                if _sfp and _utp:
                    cable = cable + " ⚠"
                uplinks.append((fd, fp, td, tp, cable))
    if not uplinks:
        return

    _add_page_break(doc)
    cell = _full_cell(doc, _rgb_hex(_C_ACCENT))
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("🔁  Uplink-overzicht")
    r1.bold = True; r1.font.size = Pt(16); r1.font.color.rgb = _C_WIT
    r2 = p.add_run(f"   —   {len(uplinks)} switch-naar-switch verbindingen")
    r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)
    _spacer(doc, 0.4)

    COLS  = ["Van switch", "Poort", "Naar switch", "Poort", "Kabeltype", "Locatie Van", "Locatie Naar"]
    W     = [Cm(3.5), Cm(2.0), Cm(3.5), Cm(2.0), Cm(2.5), Cm(6.1), Cm(8.1)]
    NCOLS = len(COLS)
    tbl   = doc.add_table(rows=0, cols=NCOLS); tbl.style = "Table Grid"

    sec_row = tbl.add_row(); _tbl_header(sec_row)
    sec_row.cells[0].merge(sec_row.cells[NCOLS - 1]); cell = sec_row.cells[0]
    _shade(cell, _rgb_hex(_C_GRIJSDK)); _no_borders(cell); _margins(cell, top=40, bottom=40, left=200, right=200)
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r = p.add_run("🔁  Switch-naar-switch uplinks"); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _C_ACCENT

    hdr_row = tbl.add_row(); _tbl_header(hdr_row)
    for i, (col, w) in enumerate(zip(COLS, W)):
        cell = hdr_row.cells[i]; cell.width = w
        _shade(cell, _rgb_hex(_C_ACCENT)); _no_borders(cell); _margins(cell, top=50, bottom=50, left=80, right=60)
        _cell_p(cell, col, bold=True, size=9, color=_C_WIT)

    for ri, (fd, fp, td, tp, cable) in enumerate(sorted(uplinks, key=lambda u: (u[0]["name"], u[1]["name"]))):
        loc_f = idx.get("loc", {}).get(fd["id"])
        loc_t = idx.get("loc", {}).get(td["id"])
        lf = f"{loc_f['room']} / {loc_f['rack']} / U{loc_f['slot']}" if loc_f else "—"
        lt = f"{loc_t['room']} / {loc_t['rack']} / U{loc_t['slot']}" if loc_t else "—"
        row = tbl.add_row(); _cant_split(row)
        bg  = _rgb_hex(_C_GRIJS) if ri % 2 == 0 else "FFFFFF"
        for i, (val, w, col) in enumerate(zip(
                [fd["name"], fp["name"], td["name"], tp["name"], cable, lf, lt], W,
                [_C_ZWART, _C_ACCENT, _C_ZWART, _C_ACCENT, _C_SUBTXT, _C_SUBTXT, _C_SUBTXT])):
            cell = row.cells[i]; cell.width = w
            _shade(cell, bg); _no_borders(cell); _margins(cell, top=35, bottom=35, left=80, right=60)
            _cell_p(cell, val, bold=(i in (0, 2)), size=9, color=col)
    _col_widths(tbl, W); _spacer(doc, 0.4)


def _add_wifi_overview(doc, data: dict, idx: dict) -> None:
    """Access Points overzicht."""
    aps = [ep for ep in data.get("endpoints", []) if ep.get("type") == "access_point"]
    if not aps:
        return
    _add_page_break(doc)
    cell = _full_cell(doc, _rgb_hex(_C_ACCENT))
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("📶  Wi-Fi — Access Points")
    r1.bold = True; r1.font.size = Pt(16); r1.font.color.rgb = _C_WIT
    r2 = p.add_run(f"   —   {len(aps)} access points")
    r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)
    _spacer(doc, 0.4)

    # F-RPT — Verbonden switch/poort per AP. Niet enkel directe endpoint<->poort
    # verbindingen, maar ook AP's die via een wandpunt door een patchpaneel-keten
    # naar een switch gepatcht zijn (_resolve_ap_uplink volgt de trace).
    ap_port_map: dict = {}
    for _ap in aps:
        _up = _resolve_ap_uplink(data, idx, _ap["id"])
        if _up:
            ap_port_map[_ap["id"]] = _up

    COLS  = ["Naam", "IP", "MAC", "Locatie", "Merk / Model", "Verbonden switch / poort"]
    W     = [Cm(4.5), Cm(3.0), Cm(4.0), Cm(4.5), Cm(4.0), Cm(7.7)]
    NCOLS = len(COLS)
    tbl   = doc.add_table(rows=0, cols=NCOLS); tbl.style = "Table Grid"

    sec_row = tbl.add_row(); _tbl_header(sec_row)
    sec_row.cells[0].merge(sec_row.cells[NCOLS - 1]); cell = sec_row.cells[0]
    _shade(cell, _rgb_hex(_C_GRIJSDK)); _no_borders(cell); _margins(cell, top=40, bottom=40, left=200, right=200)
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r = p.add_run("📶  Access Points"); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _C_ACCENT

    hdr_row = tbl.add_row(); _tbl_header(hdr_row)
    for i, (col, w) in enumerate(zip(COLS, W)):
        cell = hdr_row.cells[i]; cell.width = w
        _shade(cell, _rgb_hex(_C_ACCENT)); _no_borders(cell); _margins(cell, top=50, bottom=50, left=80, right=60)
        _cell_p(cell, col, bold=True, size=9, color=_C_WIT)

    for ri, ap in enumerate(sorted(aps, key=lambda e: e.get("name", ""))):
        port_id = ap_port_map.get(ap["id"])
        port    = idx["port"].get(port_id) if port_id else None
        dev     = idx["dev"].get(port.get("device_id", "")) if port else None
        if port and dev:
            side    = "F" if port["side"] == "front" else "B"
            sw_port = f"{dev['name']} / {port['name']} ({side}){_dev_loc_str(idx, dev['id'])}"
        else:
            sw_port = "—"
        bm  = " ".join(filter(None, [ap.get("brand", ""), ap.get("model", "")])) or "—"
        ip  = _normalize_ip(ap.get("ip", ""))
        mac = _normalize_mac(ap.get("mac", ""))
        row = tbl.add_row(); _cant_split(row)
        bg  = _rgb_hex(_C_GRIJS) if ri % 2 == 0 else "FFFFFF"
        for i, (val, w, col) in enumerate(zip(
                [ap["name"], ip, mac, ap.get("location", "") or "—", bm, sw_port], W,
                [_C_ZWART, _C_ACCENT if ip != "—" else _C_SUBTXT,
                 _C_SUBTXT, _C_SUBTXT, _C_SUBTXT, _C_ACCENT if sw_port != "—" else _C_SUBTXT])):
            cell = row.cells[i]; cell.width = w
            _shade(cell, bg); _no_borders(cell); _margins(cell, top=35, bottom=35, left=80, right=60)
            _cell_p(cell, val, bold=(i == 0), size=9, color=col)
    _col_widths(tbl, W); _spacer(doc, 0.4)


def _add_risk_outlets_section(doc, data: dict, idx: dict, risk_by_site: dict) -> None:
    """Risico wandpunten: verbonden maar zonder eindapparaat."""
    total = sum(len(v) for v in risk_by_site.values())
    if not total:
        return
    _add_page_break(doc)
    cell = _full_cell(doc, _rgb_hex(_C_ROOD))
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("⚠️  Risico wandpunten")
    r1.bold = True; r1.font.size = Pt(16); r1.font.color.rgb = _C_WIT
    r2 = p.add_run(f"   —   verbonden aan netwerk maar zonder eindapparaat  ·  {total} totaal")
    r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0xFF, 0xCC, 0xCC)
    _spacer(doc, 0.4)

    COLS  = ["Wandpunt", "Locatie", "Ruimte", "Verbonden poort", "Device", "Poort zijde"]
    W     = [Cm(2.5), Cm(4.5), Cm(4.0), Cm(3.0), Cm(6.0), Cm(7.7)]
    NCOLS = len(COLS)

    for site in get_all_sites(data):
        sn   = site["name"]
        rows = risk_by_site.get(sn, [])
        if not rows:
            continue
        cell = _full_cell(doc, _rgb_hex(_C_GRIJSDK), top=60, bottom=60)
        _clear_p(cell); p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run(f"📍  {sn}"); r1.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = _C_ACCENT
        r2 = p.add_run(f"   ·   {len(rows)} risico wandpunten"); r2.font.size = Pt(9); r2.font.color.rgb = _C_ROOD
        _spacer(doc, 0.2)

        tbl = doc.add_table(rows=0, cols=NCOLS); tbl.style = "Table Grid"
        sec_row = tbl.add_row(); _tbl_header(sec_row)
        sec_row.cells[0].merge(sec_row.cells[NCOLS - 1]); cell = sec_row.cells[0]
        _shade(cell, "FDECEA"); _no_borders(cell); _margins(cell, top=40, bottom=40, left=200, right=200)
        _clear_p(cell); p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"⚠️  Risico wandpunten  ·  {sn}"); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _C_ROOD

        hdr_row = tbl.add_row(); _tbl_header(hdr_row)
        for i, (col, w) in enumerate(zip(COLS, W)):
            cell = hdr_row.cells[i]; cell.width = w
            _shade(cell, _rgb_hex(_C_ROOD)); _no_borders(cell); _margins(cell, top=50, bottom=50, left=120, right=80)
            _cell_p(cell, col, bold=True, size=9, color=_C_WIT)

        for ri, (wo, rn, port, dev) in enumerate(sorted(rows, key=lambda x: (x[1], x[0].get("name", "")))):
            loc       = _wo_location(wo)
            port_name = port.get("name", "?") if port else "?"
            dev_name  = dev["name"] if dev else "?"
            side_str  = ("VOOR" if port["side"] == "front" else "ACHTER") if port else "?"
            row       = tbl.add_row(); _cant_split(row)
            bg        = _rgb_hex(_C_GRIJS) if ri % 2 == 0 else "FFFFFF"
            for i, (val, w, col) in enumerate(zip(
                    [wo.get("name", "?"), loc, rn, port_name, dev_name, side_str], W,
                    [_C_ZWART, _C_SUBTXT, _C_SUBTXT, _C_ROOD, _C_ACCENT, _C_SUBTXT])):
                cell = row.cells[i]; cell.width = w
                _shade(cell, bg); _no_borders(cell); _margins(cell, top=35, bottom=35, left=120, right=80)
                _cell_p(cell, val, bold=(i == 0), size=9, color=col)
        _col_widths(tbl, W); _spacer(doc, 0.4)


def _add_unwired_outlets_section(doc, data: dict, idx: dict, unwired_by_site: dict) -> None:
    """Onverbonden wandpunten: geen kabelverbinding."""
    total = sum(len(v) for v in unwired_by_site.values())
    if not total:
        return
    _add_page_break(doc)
    cell = _full_cell(doc, _rgb_hex(_C_ACCENT))
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("🔌  Onverbonden wandpunten")
    r1.bold = True; r1.font.size = Pt(16); r1.font.color.rgb = _C_WIT
    r2 = p.add_run(f"   —   geen kabelverbinding  ·  {total} totaal")
    r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)
    _spacer(doc, 0.4)

    COLS  = ["Naam", "Locatie", "Ruimte", "Eindapparaat (gekoppeld)", "Notitie"]
    W     = [Cm(2.5), Cm(5.0), Cm(4.0), Cm(6.5), Cm(9.7)]
    NCOLS = len(COLS)

    for site in get_all_sites(data):
        sn   = site["name"]
        rows = unwired_by_site.get(sn, [])
        if not rows:
            continue
        cell = _full_cell(doc, _rgb_hex(_C_GRIJSDK), top=60, bottom=60)
        _clear_p(cell); p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run(f"📍  {sn}"); r1.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = _C_ACCENT
        r2 = p.add_run(f"   ·   {len(rows)} onverbonden"); r2.font.size = Pt(9); r2.font.color.rgb = _C_SUBTXT
        _spacer(doc, 0.2)

        tbl = doc.add_table(rows=0, cols=NCOLS); tbl.style = "Table Grid"
        sec_row = tbl.add_row(); _tbl_header(sec_row)
        sec_row.cells[0].merge(sec_row.cells[NCOLS - 1]); cell = sec_row.cells[0]
        _shade(cell, _rgb_hex(_C_GRIJSDK)); _no_borders(cell); _margins(cell, top=40, bottom=40, left=200, right=200)
        _clear_p(cell); p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"🔌  Onverbonden wandpunten  ·  {sn}"); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _C_ACCENT

        hdr_row = tbl.add_row(); _tbl_header(hdr_row)
        for i, (col, w) in enumerate(zip(COLS, W)):
            cell = hdr_row.cells[i]; cell.width = w
            _shade(cell, _rgb_hex(_C_ACCENT)); _no_borders(cell); _margins(cell, top=50, bottom=50, left=120, right=80)
            _cell_p(cell, col, bold=True, size=9, color=_C_WIT)

        for ri, (wo, rn) in enumerate(sorted(rows, key=lambda x: (x[1], x[0].get("name", "")))):
            ep      = idx["ep"].get(wo.get("endpoint_id", ""))
            ep_name = ep["name"] if ep else "—"
            loc     = _wo_location(wo)
            notes   = wo.get("notes") or "—"
            row     = tbl.add_row(); _cant_split(row)
            bg      = _rgb_hex(_C_GRIJS) if ri % 2 == 0 else "FFFFFF"
            ep_col  = _C_BLAUW if ep else _C_SUBTXT
            for i, (val, w, col) in enumerate(zip(
                    [wo.get("name", "?"), loc, rn, ep_name, notes], W,
                    [_C_ZWART, _C_SUBTXT, _C_SUBTXT, ep_col, _C_SUBTXT])):
                cell = row.cells[i]; cell.width = w
                _shade(cell, bg); _no_borders(cell); _margins(cell, top=35, bottom=35, left=120, right=80)
                _cell_p(cell, val, bold=(i == 0), size=9, color=col)
        _col_widths(tbl, W); _spacer(doc, 0.4)


def _add_endpoints_section(doc, data: dict, idx: dict, ep_site_map: dict) -> None:
    """Eindapparaten — volledig overzicht per site."""
    eps = data.get("endpoints", [])
    if not eps:
        return
    _add_page_break(doc)
    cell = _full_cell(doc, _rgb_hex(_C_ACCENT))
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("🖥  Eindapparaten — volledig overzicht")
    r1.bold = True; r1.font.size = Pt(16); r1.font.color.rgb = _C_WIT
    r2 = p.add_run(f"   —   {len(eps)} totaal")
    r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)
    _spacer(doc, 0.4)

    COLS  = ["Naam", "Type", "Merk / Model", "IP", "MAC", "Locatie", "Notitie"]
    W     = [Cm(4.5), Cm(3.0), Cm(4.0), Cm(3.0), Cm(3.5), Cm(4.5), Cm(5.2)]
    NCOLS = len(COLS)

    eps_by_site: dict = {}
    for ep in eps:
        sn = ep_site_map.get(ep["id"], "Onbekend")
        eps_by_site.setdefault(sn, []).append(ep)

    for site in get_all_sites(data):
        sn   = site["name"]
        rows = eps_by_site.get(sn, [])
        if not rows:
            continue
        cell = _full_cell(doc, _rgb_hex(_C_GRIJSDK), top=60, bottom=60)
        _clear_p(cell); p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run(f"📍  {sn}"); r1.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = _C_ACCENT
        r2 = p.add_run(f"   ·   {len(rows)} eindapparaten"); r2.font.size = Pt(9); r2.font.color.rgb = _C_SUBTXT
        _spacer(doc, 0.2)

        tbl = doc.add_table(rows=0, cols=NCOLS); tbl.style = "Table Grid"
        sec_row = tbl.add_row(); _tbl_header(sec_row)
        sec_row.cells[0].merge(sec_row.cells[NCOLS - 1]); cell = sec_row.cells[0]
        _shade(cell, _rgb_hex(_C_GRIJSDK)); _no_borders(cell); _margins(cell, top=40, bottom=40, left=200, right=200)
        _clear_p(cell); p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"🖥  Eindapparaten  ·  {sn}"); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _C_ACCENT

        hdr_row = tbl.add_row(); _tbl_header(hdr_row)
        for i, (col, w) in enumerate(zip(COLS, W)):
            cell = hdr_row.cells[i]; cell.width = w
            _shade(cell, _rgb_hex(_C_ACCENT)); _no_borders(cell); _margins(cell, top=50, bottom=50, left=120, right=80)
            _cell_p(cell, col, bold=True, size=9, color=_C_WIT)

        for ri, ep in enumerate(sorted(rows, key=lambda e: e.get("name", ""))):
            ep_type = _EP_TYPE_MAP.get(ep.get("type", ""), ep.get("type", "") or "—")
            bm      = " ".join(filter(None, [ep.get("brand", ""), ep.get("model", "")])) or "—"
            ip      = _normalize_ip(ep.get("ip", ""))
            mac     = _normalize_mac(ep.get("mac", ""))
            loc     = ep.get("location", "") or "—"
            notes   = ep.get("notes", "") or "—"
            row     = tbl.add_row(); _cant_split(row)
            bg      = _rgb_hex(_C_GRIJS) if ri % 2 == 0 else "FFFFFF"
            for i, (val, w, col) in enumerate(zip(
                    [ep["name"], ep_type, bm, ip, mac, loc, notes], W,
                    [_C_ZWART, _C_SUBTXT, _C_SUBTXT,
                     _C_ACCENT if ip != "—" else _C_SUBTXT,
                     _C_SUBTXT, _C_SUBTXT, _C_SUBTXT])):
                cell = row.cells[i]; cell.width = w
                _shade(cell, bg); _no_borders(cell); _margins(cell, top=35, bottom=35, left=120, right=80)
                _cell_p(cell, val, bold=(i == 0), size=9, color=col)
        _col_widths(tbl, W); _spacer(doc, 0.4)

    unknown = eps_by_site.get("Onbekend", [])
    if unknown:
        p = doc.add_paragraph()
        r = p.add_run(f"  ℹ️  {len(unknown)} eindapparaten zonder sitekoppeling (niet verbonden via poort of wandpunt)")
        r.font.size = Pt(9); r.font.color.rgb = _C_SUBTXT; r.italic = True


def _add_device_info_section(doc, data: dict, idx: dict, dup_ips: set) -> None:
    """Device info uitgebreid: MAC (genorm.), IP, serial — met dubbele-IP markering."""
    devs_with_info = [dv for dv in data.get("devices", []) if dv.get("mac") or dv.get("serial") or dv.get("ip")]
    if not devs_with_info:
        return
    _add_page_break(doc)
    cell = _full_cell(doc, _rgb_hex(_C_ACCENT))
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("🗄  Device info — uitgebreid")
    r1.bold = True; r1.font.size = Pt(16); r1.font.color.rgb = _C_WIT
    r2 = p.add_run(f"   —   {len(devs_with_info)} devices met MAC / serial / IP")
    r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)
    _spacer(doc, 0.4)

    COLS  = ["Device", "Type", "IP", "MAC (genorm.)", "Serial", "Locatie"]
    W     = [Cm(4.5), Cm(3.0), Cm(3.5), Cm(4.5), Cm(4.5), Cm(7.7)]
    NCOLS = len(COLS)

    for site in get_all_sites(data):
        site_dev_ids = {
            slot.get("device_id")
            for room in site.get("rooms", [])
            for rack in room.get("racks", [])
            for slot in rack.get("slots", [])
            if slot.get("device_id")
        }
        site_devs = [dv for dv in devs_with_info if dv["id"] in site_dev_ids]
        if not site_devs:
            continue
        cell = _full_cell(doc, _rgb_hex(_C_GRIJSDK), top=60, bottom=60)
        _clear_p(cell); p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run(f"📍  {site['name']}"); r1.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = _C_ACCENT
        r2 = p.add_run(f"   ·   {len(site_devs)} devices"); r2.font.size = Pt(9); r2.font.color.rgb = _C_SUBTXT
        _spacer(doc, 0.2)

        tbl = doc.add_table(rows=0, cols=NCOLS); tbl.style = "Table Grid"
        sec_row = tbl.add_row(); _tbl_header(sec_row)
        sec_row.cells[0].merge(sec_row.cells[NCOLS - 1]); cell = sec_row.cells[0]
        _shade(cell, _rgb_hex(_C_GRIJSDK)); _no_borders(cell); _margins(cell, top=40, bottom=40, left=200, right=200)
        _clear_p(cell); p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"🗄  Device info  ·  {site['name']}"); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _C_ACCENT

        hdr_row = tbl.add_row(); _tbl_header(hdr_row)
        for i, (col, w) in enumerate(zip(COLS, W)):
            cell = hdr_row.cells[i]; cell.width = w
            _shade(cell, _rgb_hex(_C_ACCENT)); _no_borders(cell); _margins(cell, top=50, bottom=50, left=120, right=80)
            _cell_p(cell, col, bold=True, size=9, color=_C_WIT)

        for ri, dv in enumerate(sorted(site_devs, key=lambda d: d.get("name", ""))):
            tl     = _TYPE_MAP.get(dv.get("type", ""), dv.get("type", "") or "—")
            ip     = _normalize_ip(dv.get("ip", ""))
            mac    = _normalize_mac(dv.get("mac", ""))
            serial = dv.get("serial", "") or "—"
            loc    = idx.get("loc", {}).get(dv["id"])
            loc_str = f"{loc['room']} / {loc['rack']} / U{loc['slot']}" if loc else "—"
            ip_dup = ip in dup_ips and ip != "—"
            row    = tbl.add_row(); _cant_split(row)
            bg     = _rgb_hex(_C_GRIJS) if ri % 2 == 0 else "FFFFFF"
            ip_col = _C_ORANJE if ip_dup else (_C_ACCENT if ip != "—" else _C_SUBTXT)
            for i, (val, w, col) in enumerate(zip(
                    [dv["name"], tl, ip + (" ⚠" if ip_dup else ""), mac, serial, loc_str], W,
                    [_C_ZWART, _C_SUBTXT, ip_col, _C_SUBTXT, _C_SUBTXT, _C_SUBTXT])):
                cell = row.cells[i]; cell.width = w
                _shade(cell, bg); _no_borders(cell); _margins(cell, top=35, bottom=35, left=120, right=80)
                _cell_p(cell, val, bold=(i == 0), size=9, color=col)
        _col_widths(tbl, W); _spacer(doc, 0.4)


def _compute_action_items(data: dict, idx: dict,
                          risk_by_site: dict, unwired_by_site: dict,
                          dup_ips: set) -> list:
    """Bouwt de actieplan-lijst (zonder renderen). Telt vanuit de OPEN
    actie-objecten — goedgekeurde uitzonderingen (data["approvals"]) zijn
    afgetrokken. Eén bron (_enumerate_action_objects) voor cover-telling,
    actieplan-tabel en reviewtool. Drempelcheck VLAN-info blijft aggregaat.
    Elk item: (nr, bevinding, toelichting, prioriteit, actie, status)."""
    from collections import defaultdict

    approved  = _approved_keys(data)
    objs      = _enumerate_action_objects(data, idx, risk_by_site, unwired_by_site, dup_ips)
    open_objs = [o for o in objs if o["key"] not in approved]
    by: dict = defaultdict(list)
    for o in open_objs:
        by[o["check"]].append(o)

    acties = []
    nr = 1

    def _add(bevinding, toelichting, prio, actie):
        nonlocal nr
        acties.append((nr, bevinding, toelichting, prio, actie, "Open"))
        nr += 1

    if by["risk_outlet"]:
        _add(f"Risico wandpunten ({len(by['risk_outlet'])})",
             "Verbonden aan netwerk zonder eindapparaat — onbeheerde netwerktoegang",
             "Hoog", "Fysiek controleren en eindapparaat registreren of poort blokkeren")
    if by["unwired_outlet"]:
        _add(f"Onverbonden wandpunten ({len(by['unwired_outlet'])})",
             "Geen kabelverbinding geregistreerd — onduidelijk of fysiek aangesloten",
             "Middel", "Fysiek controleren en verbinding registreren of wandpunt verwijderen")
    if by["dup_ip"]:
        _ips = sorted(o["object_id"] for o in by["dup_ip"])
        _add(f"Dubbele IP-adressen ({len(_ips)})",
             f"Meerdere devices delen zelfde IP: {', '.join(_ips)}",
             "Hoog", "IP-conflict corrigeren (switches met # in naam en zelfde prefix zijn al uitgesloten)")
    if by["unknown_ep"]:
        _add(f"Eindapparaten zonder sitekoppeling ({len(by['unknown_ep'])})",
             "Niet verbonden via poort of wandpunt",
             "Laag", "Verbinding registreren of eindapparaat verwijderen")

    # Drempelcheck — aggregaat, niet per-object goedkeurbaar
    ports_no_vlan = sum(1 for p in data.get("ports", []) if not p.get("vlan") and p.get("side") == "front")
    if ports_no_vlan > 50:
        _add(f"Poorten zonder VLAN-info ({ports_no_vlan})",
             "VLAN-kolom leeg — onduidelijk of access, trunk of niet in gebruik",
             "Middel", "VLAN-informatie aanvullen per switchpoort")

    if by["bad_mac"]:
        _names = [o["_name"] for o in by["bad_mac"]]
        _suffix = "..." if len(_names) > 5 else ""
        _add(f"MAC-adressen in brondata niet genormaliseerd ({len(_names)})",
             f"Brondata bevat MAC zonder scheidingstekens: {', '.join(_names[:5])}{_suffix}. "
             "Het rapport normaliseert visueel, brondata moet nog gecorrigeerd worden.",
             "Laag", "Brondata aanpassen naar AA:BB:CC:DD:EE:FF formaat")
    if by["no_brand_model"]:
        _names = [o["_name"] for o in by["no_brand_model"]]
        _add(f"Switches zonder merk/model ({len(_names)})",
             f"{', '.join(_names[:5])}",
             "Laag", "Merk en model aanvullen voor correcte inventaris")
    if by["sfp_utp"]:
        _items = [o["_desc"] for o in by["sfp_utp"]]
        _add(f"SFP-uplinks met mogelijk verkeerd kabeltype ({len(_items)})",
             f"{'; '.join(_items[:3])}{'...' if len(_items) > 3 else ''}",
             "Middel", "Kabeltype controleren: fiber, DAC of koper-SFP")

    # Wi-Fi access points — datakwaliteit (gecombineerde lijn, open-telling)
    _ap_no_ip   = len(by["ap_no_ip"])
    _ap_no_port = len(by["ap_no_uplink"])
    _ap_dup     = len({o["object_id"] for o in by["ap_dup_mac"]})
    _ap_det = []
    if _ap_no_ip:   _ap_det.append(f"{_ap_no_ip} zonder IP")
    if _ap_no_port: _ap_det.append(f"{_ap_no_port} zonder switch/poort")
    if _ap_dup:     _ap_det.append(f"{_ap_dup} dubbele identiteit (MAC)")
    if _ap_det:
        _add("Wi-Fi access points — datakwaliteit",
             "; ".join(_ap_det),
             "Middel", "IP, poort en dubbele entries aanvullen/corrigeren")

    return acties


def _add_action_plan(doc, data: dict, idx: dict,
                     risk_by_site: dict, unwired_by_site: dict, dup_ips: set) -> None:
    """Auto-gegenereerd actieplan op basis van data-kwaliteitschecks."""
    _add_page_break(doc)
    cell = _full_cell(doc, _rgb_hex(_C_ACCENT))
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r = p.add_run("📋  Actieplan — open punten")
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = _C_WIT
    _spacer(doc, 0.4)

    acties = _compute_action_items(data, idx, risk_by_site, unwired_by_site, dup_ips)
    # Issue-aantallen opslaan voor cover
    data["_action_counts"] = {
        "total":  len(acties),
        "high":   sum(1 for a in acties if a[3] == "Hoog"),
        "medium": sum(1 for a in acties if a[3] == "Middel"),
        "low":    sum(1 for a in acties if a[3] == "Laag"),
    }
    PRIO_COLOR = {"Hoog": _C_ROOD, "Middel": _C_ORANJE, "Laag": _C_GROEN}
    COLS  = ["#", "Bevinding", "Toelichting", "Prioriteit", "Actie", "Status"]
    W     = [Cm(0.8), Cm(4.5), Cm(7.0), Cm(2.0), Cm(8.0), Cm(5.4)]
    NCOLS = len(COLS)
    tbl   = doc.add_table(rows=0, cols=NCOLS); tbl.style = "Table Grid"

    sec_row = tbl.add_row(); _tbl_header(sec_row)
    sec_row.cells[0].merge(sec_row.cells[NCOLS - 1]); cell = sec_row.cells[0]
    _shade(cell, _rgb_hex(_C_GRIJSDK)); _no_borders(cell); _margins(cell, top=40, bottom=40, left=200, right=200)
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"📋  {len(acties)} open actiepunten"); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _C_ACCENT

    hdr_row = tbl.add_row(); _tbl_header(hdr_row)
    for i, (col, w) in enumerate(zip(COLS, W)):
        cell = hdr_row.cells[i]; cell.width = w
        _shade(cell, _rgb_hex(_C_ACCENT)); _no_borders(cell); _margins(cell, top=50, bottom=50, left=80, right=60)
        _cell_p(cell, col, bold=True, size=9, color=_C_WIT)

    for ri, (num, bevinding, toelichting, prio, actie, status) in enumerate(acties):
        row = tbl.add_row(); _cant_split(row)
        bg  = _rgb_hex(_C_GRIJS) if ri % 2 == 0 else "FFFFFF"
        pc  = PRIO_COLOR.get(prio, _C_SUBTXT)
        for i, (val, w, col) in enumerate(zip(
                [str(num), bevinding, toelichting, prio, actie, status], W,
                [_C_SUBTXT, _C_ZWART, _C_SUBTXT, pc, _C_SUBTXT, _C_SUBTXT])):
            cell = row.cells[i]; cell.width = w
            _shade(cell, bg); _no_borders(cell); _margins(cell, top=35, bottom=35, left=80, right=60)
            _cell_p(cell, val, bold=(i == 1), size=9, color=col)
    _col_widths(tbl, W); _spacer(doc, 0.4)


def _add_approved_exceptions(doc, approved_objs: list) -> None:
    """Goedgekeurde uitzonderingen: actiepunten die manueel zijn goedgekeurd en
    daardoor NIET meetellen in het actieplan. Jouw keuze: apart tonen."""
    if not approved_objs:
        return
    _add_page_break(doc)
    cell = _full_cell(doc, _rgb_hex(_C_GROEN))
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("\u2705  Goedgekeurde uitzonderingen")
    r1.bold = True; r1.font.size = Pt(16); r1.font.color.rgb = _C_WIT
    r2 = p.add_run(f"   \u2014   {len(approved_objs)} handmatig goedgekeurd \u00b7 niet meegeteld in het actieplan")
    r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0xDD, 0xEE, 0xDD)
    _spacer(doc, 0.4)

    COLS  = ["Categorie", "Object", "Bevinding", "Reden", "Door", "Datum"]
    W     = [Cm(4.0), Cm(5.0), Cm(6.0), Cm(8.0), Cm(3.0), Cm(3.7)]
    NCOLS = len(COLS)
    tbl   = doc.add_table(rows=0, cols=NCOLS); tbl.style = "Table Grid"

    hdr_row = tbl.add_row(); _tbl_header(hdr_row)
    for i, (col, w) in enumerate(zip(COLS, W)):
        cell = hdr_row.cells[i]; cell.width = w
        _shade(cell, _rgb_hex(_C_GROEN)); _no_borders(cell); _margins(cell, top=50, bottom=50, left=80, right=60)
        _cell_p(cell, col, bold=True, size=9, color=_C_WIT)

    for ri, o in enumerate(sorted(approved_objs, key=lambda x: (x["category"], x["label"]))):
        ap  = o.get("_approval", {})
        at  = (ap.get("at", "") or "")[:10]
        row = tbl.add_row(); _cant_split(row)
        bg  = _rgb_hex(_C_GRIJS) if ri % 2 == 0 else "FFFFFF"
        vals = [o["category"], o["label"], o["detail"],
                ap.get("reason", "") or "\u2014",
                ap.get("by", "") or "\u2014", at or "\u2014"]
        for i, (val, w) in enumerate(zip(vals, W)):
            cell = row.cells[i]; cell.width = w
            _shade(cell, bg); _no_borders(cell); _margins(cell, top=35, bottom=35, left=80, right=60)
            _cell_p(cell, val, bold=(i == 1), size=9,
                    color=(_C_ZWART if i == 1 else _C_SUBTXT))
    _col_widths(tbl, W); _spacer(doc, 0.4)


def _add_orphans_section(doc, data: dict, idx: dict) -> None:
    """Niet-geplaatste objecten: devices zonder rack-slot en eindapparaten zonder
    site-/poortkoppeling. Rendert enkel als zulke objecten bestaan — in een
    bedrijfsgefilterd rapport is dit normaal leeg."""
    placed = {
        slot.get("device_id")
        for site in get_all_sites(data)
        for room in site.get("rooms", [])
        for rack in room.get("racks", [])
        for slot in rack.get("slots", [])
        if slot.get("device_id")
    }
    orphan_devs = [dv for dv in data.get("devices", []) if dv["id"] not in placed]
    ep_map      = _build_ep_site_map(data, idx)
    orphan_eps  = [ep for ep in data.get("endpoints", []) if ep["id"] not in ep_map]
    if not orphan_devs and not orphan_eps:
        return

    _add_page_break(doc)
    cell = _full_cell(doc, _rgb_hex(_C_ACCENT))
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("\U0001F9E9  Niet-geplaatste objecten")
    r1.bold = True; r1.font.size = Pt(16); r1.font.color.rgb = _C_WIT
    r2 = p.add_run(f"   \u2014   {len(orphan_devs)} devices  \u00b7  {len(orphan_eps)} eindapparaten")
    r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)
    _spacer(doc, 0.3)

    note = doc.add_paragraph()
    rn = note.add_run(
        "  \u2139\ufe0f  Deze objecten staan in de databron maar zijn niet aan een rack-slot "
        "of verbinding gekoppeld. Ze tellen mee in de totalen op het titelblad, maar verschijnen "
        "niet in de site-/rackdetails en evenmin in een bedrijfsgefilterd rapport."
    )
    rn.font.size = Pt(9); rn.italic = True; rn.font.color.rgb = _C_SUBTXT
    _spacer(doc, 0.2)

    if orphan_devs:
        COLS = ["Device", "Type", "Merk / Model", "IP", "MAC (genorm.)", "Serial"]
        W    = [Cm(5.0), Cm(3.5), Cm(5.5), Cm(3.5), Cm(4.5), Cm(5.7)]
        NCOLS = len(COLS)
        tbl = doc.add_table(rows=0, cols=NCOLS); tbl.style = "Table Grid"
        sec_row = tbl.add_row(); _tbl_header(sec_row)
        sec_row.cells[0].merge(sec_row.cells[NCOLS - 1]); cell = sec_row.cells[0]
        _shade(cell, _rgb_hex(_C_GRIJSDK)); _no_borders(cell); _margins(cell, top=40, bottom=40, left=200, right=200)
        _clear_p(cell); p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"\U0001F5C4  Devices zonder rack-plaatsing  \u00b7  {len(orphan_devs)}")
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _C_ACCENT
        hdr_row = tbl.add_row(); _tbl_header(hdr_row)
        for i, (col, w) in enumerate(zip(COLS, W)):
            cell = hdr_row.cells[i]; cell.width = w
            _shade(cell, _rgb_hex(_C_ACCENT)); _no_borders(cell); _margins(cell, top=50, bottom=50, left=120, right=80)
            _cell_p(cell, col, bold=True, size=9, color=_C_WIT)
        for ri, dv in enumerate(sorted(orphan_devs, key=lambda d: d.get("name", ""))):
            tl  = _TYPE_MAP.get(dv.get("type", ""), dv.get("type", "") or "\u2014")
            bm  = " ".join(filter(None, [dv.get("brand", ""), dv.get("model", "")])) or "\u2014"
            ip  = _normalize_ip(dv.get("ip", ""))
            mac = _normalize_mac(dv.get("mac", ""))
            ser = dv.get("serial", "") or "\u2014"
            row = tbl.add_row(); _cant_split(row)
            bg  = _rgb_hex(_C_GRIJS) if ri % 2 == 0 else "FFFFFF"
            for i, (val, w, col) in enumerate(zip(
                    [dv.get("name", "\u2014"), tl, bm, ip, mac, ser], W,
                    [_C_ZWART, _C_SUBTXT, _C_SUBTXT,
                     _C_ACCENT if ip != "\u2014" else _C_SUBTXT, _C_SUBTXT, _C_SUBTXT])):
                cell = row.cells[i]; cell.width = w
                _shade(cell, bg); _no_borders(cell); _margins(cell, top=35, bottom=35, left=120, right=80)
                _cell_p(cell, val, bold=(i == 0), size=9, color=col)
        _col_widths(tbl, W); _spacer(doc, 0.3)

    if orphan_eps:
        COLS = ["Eindapparaat", "Type", "Merk / Model", "IP", "MAC (genorm.)", "Notitie"]
        W    = [Cm(5.0), Cm(3.5), Cm(5.5), Cm(3.5), Cm(4.5), Cm(5.7)]
        NCOLS = len(COLS)
        tbl = doc.add_table(rows=0, cols=NCOLS); tbl.style = "Table Grid"
        sec_row = tbl.add_row(); _tbl_header(sec_row)
        sec_row.cells[0].merge(sec_row.cells[NCOLS - 1]); cell = sec_row.cells[0]
        _shade(cell, _rgb_hex(_C_GRIJSDK)); _no_borders(cell); _margins(cell, top=40, bottom=40, left=200, right=200)
        _clear_p(cell); p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"\U0001F5A5  Eindapparaten zonder koppeling  \u00b7  {len(orphan_eps)}")
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _C_ACCENT
        hdr_row = tbl.add_row(); _tbl_header(hdr_row)
        for i, (col, w) in enumerate(zip(COLS, W)):
            cell = hdr_row.cells[i]; cell.width = w
            _shade(cell, _rgb_hex(_C_ACCENT)); _no_borders(cell); _margins(cell, top=50, bottom=50, left=120, right=80)
            _cell_p(cell, col, bold=True, size=9, color=_C_WIT)
        for ri, ep in enumerate(sorted(orphan_eps, key=lambda e: e.get("name", ""))):
            tl   = _EP_TYPE_MAP.get(ep.get("type", ""), ep.get("type", "") or "\u2014")
            bm   = " ".join(filter(None, [ep.get("brand", ""), ep.get("model", "")])) or "\u2014"
            ip   = _normalize_ip(ep.get("ip", ""))
            mac  = _normalize_mac(ep.get("mac", ""))
            note = ep.get("notes", "") or "\u2014"
            row  = tbl.add_row(); _cant_split(row)
            bg   = _rgb_hex(_C_GRIJS) if ri % 2 == 0 else "FFFFFF"
            for i, (val, w, col) in enumerate(zip(
                    [ep.get("name", "\u2014"), tl, bm, ip, mac, note], W,
                    [_C_ZWART, _C_SUBTXT, _C_SUBTXT,
                     _C_ACCENT if ip != "\u2014" else _C_SUBTXT, _C_SUBTXT, _C_SUBTXT])):
                cell = row.cells[i]; cell.width = w
                _shade(cell, bg); _no_borders(cell); _margins(cell, top=35, bottom=35, left=120, right=80)
                _cell_p(cell, val, bold=(i == 0), size=9, color=col)
        _col_widths(tbl, W); _spacer(doc, 0.3)


def _add_revision_history(doc, version: str) -> None:
    """Revisiehistoriek — eerste rij ingevuld, rest blanco voor opvolging."""
    _add_page_break(doc)
    cell = _full_cell(doc, _rgb_hex(_C_ACCENT))
    _clear_p(cell); p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    r = p.add_run("📝  Revisiehistoriek"); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = _C_WIT
    _spacer(doc, 0.4)

    COLS  = ["Versie", "Datum", "Auteur", "Omschrijving"]
    W     = [Cm(2.0), Cm(3.0), Cm(4.0), Cm(18.7)]
    tbl   = doc.add_table(rows=0, cols=4); tbl.style = "Table Grid"
    hdr_row = tbl.add_row(); _tbl_header(hdr_row)
    for i, (col, w) in enumerate(zip(COLS, W)):
        cell = hdr_row.cells[i]; cell.width = w
        _shade(cell, _rgb_hex(_C_ACCENT)); _no_borders(cell); _margins(cell, top=50, bottom=50, left=120, right=80)
        _cell_p(cell, col, bold=True, size=9, color=_C_WIT)

    revisions = [
        (version, datetime.date.today().strftime("%d/%m/%Y"), "Networkmap Creator",
         "Initiële export — automatisch gegenereerd rapport"),
        ("", "", "", ""),
        ("", "", "", ""),
    ]
    for ri, (ver, dat, aut, omschr) in enumerate(revisions):
        row = tbl.add_row(); _cant_split(row)
        bg  = _rgb_hex(_C_GRIJS) if ri % 2 == 0 else "FFFFFF"
        for i, (val, w, col) in enumerate(zip(
                [ver, dat, aut, omschr], W,
                [_C_ACCENT, _C_SUBTXT, _C_SUBTXT, _C_ZWART])):
            cell = row.cells[i]; cell.width = w
            _shade(cell, bg); _no_borders(cell); _margins(cell, top=40, bottom=40, left=120, right=80)
            _cell_p(cell, val, bold=(i == 0 and bool(val)), size=9, color=col)
    _col_widths(tbl, W)


# ===========================================================================
# HOOFD RENDER FUNCTIE
# ===========================================================================

def _scope_data_to_company(data: dict, company: dict) -> dict:
    """
    F3 — Geeft een gefilterde kopie van data met enkel objecten van één bedrijf.

    companies[] wordt beperkt tot dit ene bedrijf; devices, ports, endpoints en
    connections worden gefilterd op membership binnen de sites van dat bedrijf.
    De oorspronkelijke data dict blijft ongewijzigd (shallow copy + nieuwe lijsten).
    """
    dev_ids: set = set()
    outlet_ids: set = set()
    for site in company.get("sites", []):
        for room in site.get("rooms", []):
            for rack in room.get("racks", []):
                for slot in rack.get("slots", []):
                    did = slot.get("device_id")
                    if did:
                        dev_ids.add(did)
            for wo in room.get("wall_outlets", []):
                if wo.get("id"):
                    outlet_ids.add(wo["id"])

    ports    = [p for p in data.get("ports", []) if p.get("device_id") in dev_ids]
    port_ids = {p["id"] for p in ports}
    devices  = [d for d in data.get("devices", []) if d.get("id") in dev_ids]

    ep_ids: set = set()
    for site in company.get("sites", []):
        for room in site.get("rooms", []):
            for wo in room.get("wall_outlets", []):
                eid = wo.get("endpoint_id")
                if eid:
                    ep_ids.add(eid)

    def _phys_in(end_type, end_id) -> bool:
        if end_type == "port":
            return end_id in port_ids
        if end_type == "wall_outlet":
            return end_id in outlet_ids
        return False

    conns: list = []
    for c in data.get("connections", []):
        ft, fi = c.get("from_type"), c.get("from_id")
        tt, ti = c.get("to_type"),   c.get("to_id")
        if _phys_in(ft, fi) or _phys_in(tt, ti):
            conns.append(c)
            if ft == "endpoint" and fi:
                ep_ids.add(fi)
            if tt == "endpoint" and ti:
                ep_ids.add(ti)

    endpoints = [e for e in data.get("endpoints", []) if e.get("id") in ep_ids]

    scoped = dict(data)
    scoped["companies"]   = [company]
    scoped["devices"]     = devices
    scoped["ports"]       = ports
    scoped["endpoints"]   = endpoints
    scoped["connections"] = conns
    return scoped


# ===========================================================================
# RW-rapport — bedrijfsgroepering: koppen, inhoudsopgave, meta
# ===========================================================================

def _company_heading(doc, name: str, is_first: bool) -> None:
    """Heading 1 per bedrijf (TOC-entry). Paginabreuk vooraf, behalve het eerste."""
    h = doc.add_paragraph(name, style="Heading 1")
    if not is_first:
        h.paragraph_format.page_break_before = True


def _rubriek(doc, title: str, page_break: bool = True) -> None:
    """Heading 2 rubriektitel (TOC-entry).

    Plaatst optioneel een paginabreuk vóór de rubriek en onderdrukt de eigen
    paginabreuk van de daaropvolgende sectie, zodat de gekleurde sectiebanner op
    dezelfde pagina komt als deze kop (geen lege tussenpagina).
    """
    global _SUPPRESS_NEXT_BREAK
    h = doc.add_paragraph(title, style="Heading 2")
    if page_break:
        h.paragraph_format.page_break_before = True
    _SUPPRESS_NEXT_BREAK = True


def _add_toc(doc) -> None:
    """Voegt een automatisch bijwerkbaar Word-inhoudsopgaveveld toe (Heading 1-3).

    Word vult het veld bij het openen van het document (of via rechtsklik ->
    'Veld bijwerken'). De bedrijven (Heading 1) en rubrieken (Heading 2)
    verschijnen dan automatisch met paginanummers.
    """
    title = doc.add_paragraph()
    rt = title.add_run("Inhoudsopgave")
    rt.bold = True
    rt.font.size = Pt(18)
    rt.font.color.rgb = _C_ACCENT
    _spacer(doc, 0.2)

    p = doc.add_paragraph()
    run = p.add_run()
    f_begin = OxmlElement("w:fldChar"); f_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    f_sep = OxmlElement("w:fldChar"); f_sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "Rechtsklik en kies 'Veld bijwerken' om de inhoudsopgave te genereren."
    f_end = OxmlElement("w:fldChar"); f_end.set(qn("w:fldCharType"), "end")
    run._r.append(f_begin); run._r.append(instr); run._r.append(f_sep)
    run._r.append(t); run._r.append(f_end)


def _has_uplinks(data: dict, idx: dict) -> bool:
    """True als er minstens één switch-naar-switch uplink bestaat (zelfde detectie
    als _add_uplink_overview), zodat geen lege rubriekkop wordt geplaatst."""
    switch_ids = {dv["id"] for dv in data.get("devices", []) if dv.get("type") == "switch"}
    if len(switch_ids) < 2:
        return False
    for conn in data.get("connections", []):
        fp = idx["port"].get(conn.get("from_id"))
        tp = idx["port"].get(conn.get("to_id"))
        if not fp or not tp:
            continue
        fd = idx["dev"].get(fp.get("device_id", ""))
        td = idx["dev"].get(tp.get("device_id", ""))
        if (fd and td and fd["id"] in switch_ids and td["id"] in switch_ids
                and fd["id"] != td["id"]):
            return True
    return False


def _prepare_report_meta(data: dict, idx: dict, risk_by_site: dict,
                         unwired_by_site: dict, dup_ips: set,
                         vlan_config: dict) -> None:
    """Berekent VLAN-tellingen, validatiestatus en actieplan-aantallen en zet ze
    op de meegegeven data dict (_vlan_*, _validation_status, _action_counts,
    _vlan_config). Wordt zowel globaal (cover) als per bedrijf gebruikt zodat de
    tellingen telkens kloppen voor de betreffende scope."""
    sites = get_all_sites(data)

    all_vlans_used: set = set()
    for p in data.get("ports", []):
        if p.get("vlan"):
            try: all_vlans_used.add(int(p["vlan"]))
            except Exception: all_vlans_used.add(p["vlan"])
    for s in sites:
        for room in s.get("rooms", []):
            for wo in room.get("wall_outlets", []):
                if wo.get("vlan"):
                    try: all_vlans_used.add(int(wo["vlan"]))
                    except Exception: all_vlans_used.add(wo["vlan"])
    data["_vlan_count_corrected"] = len(all_vlans_used)
    data["_vlan_count_used"]      = len(all_vlans_used)
    data["_vlan_config"]          = vlan_config

    if vlan_config:
        import ipaddress as _ipa
        _n_def = len(vlan_config)
        _n_complete = 0
        _n_issues = 0
        _seen_gw: dict = {}
        for _vid, _vcfg in vlan_config.items():
            _ip  = _vcfg.get("ip", "")
            _sub = _vcfg.get("subnet", "")
            _ok_ip, _ok_sub = False, False
            if _ip:
                try: _ipa.ip_address(_ip); _ok_ip = True
                except Exception: _n_issues += 1
                if _ok_ip:
                    _seen_gw.setdefault(_ip, []).append(_vid)
            else:
                _n_issues += 1
            if _sub:
                try: _ipa.ip_network(f"0.0.0.0/{_sub}", strict=False); _ok_sub = True
                except Exception: _n_issues += 1
            if _ok_ip and _ok_sub and _vcfg.get("name"):
                _n_complete += 1
        for _ip, _vids in _seen_gw.items():
            if len(_vids) > 1:
                _n_issues += 1
        data["_vlan_count_defined"]  = _n_def
        data["_vlan_count_complete"] = _n_complete
        data["_vlan_issues"]         = _n_issues if _n_issues else "—"
    else:
        data["_vlan_count_defined"]  = 0
        data["_vlan_count_complete"] = 0
        data["_vlan_issues"]         = "—"

    _risk_total = sum(len(v) for v in risk_by_site.values())
    _high   = bool(dup_ips) or (_risk_total > 0)
    _medium = bool(unwired_by_site and any(unwired_by_site.values()))
    if _high:
        data["_validation_status"] = "NIET DEFINITIEF"
    elif _medium:
        data["_validation_status"] = "TE CONTROLEREN"
    else:
        data["_validation_status"] = "GEVALIDEERD"

    _ap_items = _compute_action_items(data, idx, risk_by_site, unwired_by_site, dup_ips)
    data["_action_counts"] = {
        "total":  len(_ap_items),
        "high":   sum(1 for a in _ap_items if a[3] == "Hoog"),
        "medium": sum(1 for a in _ap_items if a[3] == "Middel"),
        "low":    sum(1 for a in _ap_items if a[3] == "Laag"),
    }


def render_report_docx(data: dict, filepath: str, company_id: str = "") -> tuple[bool, str]:
    """
    Genereer volledig Word rapport met bedrijfsgroepering (RW-rapport).

    company_id == ""  -> render alle bedrijven. Elk bedrijf krijgt een eigen
                         Heading 1-sectie (paginabreuk vooraf, behalve het eerste)
                         met al zijn rubrieken; data wordt per bedrijf gefilterd via
                         _scope_data_to_company(). Een automatische Word-TOC toont de
                         bedrijven en rubrieken.
    company_id != ""  -> enkel dat ene bedrijf (structuur identiek, één blok).

    Achterwaarts compatibel: bestaande aanroepen zonder company_id blijven werken.
    Returns: (success: bool, error_message: str)
    """
    try:
        from app import version as _ver
        version = _ver.__version__
    except Exception:
        version = "2.0.0"

    # Bedrijfsselectie / -scope
    if company_id:
        sel = get_company_by_id(data, company_id)
        if sel is None:
            return False, f"Bedrijf niet gevonden: {company_id}"
        companies = [sel]
    else:
        companies = list(get_all_companies(data))
    if not companies:
        return False, "Geen bedrijven in de data."

    try:
        doc = Document()
        sec = doc.sections[0]
        sec.orientation   = WD_ORIENT.LANDSCAPE
        sec.page_width    = Cm(29.7)
        sec.page_height   = Cm(21.0)
        sec.top_margin    = Cm(1.8)
        sec.bottom_margin = Cm(1.2)
        sec.left_margin   = Cm(1.5)
        sec.right_margin  = Cm(1.5)

        # VLAN-configuratie (globaal, niet bedrijfsspecifiek) — eenmaal laden
        vlan_config: dict = {}
        try:
            from app.services import vlan_service
            vlans_raw = vlan_service.load_vlans()
            vlan_config = {v["id"]: v for v in vlans_raw}
        except Exception:
            pass

        # Cover-scope: bij één bedrijf gefilterd, anders volledige data (alle bedrijven)
        full_data = _scope_data_to_company(data, companies[0]) if company_id else data
        g_idx   = _build_index(full_data)
        g_sites = get_all_sites(full_data)

        cover_company = companies[0] if len(companies) == 1 else None
        if cover_company is not None:
            org_name = cover_company.get("name") or (g_sites[0]["name"] if g_sites else "Netwerk")
        else:
            org_name = " · ".join(c.get("name", "") for c in companies)

        full_data["_company"] = cover_company
        g_risk, g_unwired = _build_wo_risk_maps(full_data, g_idx)
        g_dup             = _build_duplicate_ip_set(full_data)
        g_risk, g_unwired = _strip_approved_outlets(g_risk, g_unwired, _approved_keys(full_data))
        _prepare_report_meta(full_data, g_idx, g_risk, g_unwired, g_dup, vlan_config)

        _add_header(doc, org_name, version)
        _add_footer(doc, version)

        # 1. Titelblad (globaal) + automatische inhoudsopgave
        _build_titlepage(doc, full_data, version)
        _add_page_break(doc)
        _add_toc(doc)
        _add_page_break(doc)

        # 2. Per bedrijf: alle rubrieken gefilterd op dat bedrijf
        for i, comp in enumerate(companies):
            sdata = full_data if company_id else _scope_data_to_company(data, comp)
            sdata["_company"] = comp
            sidx          = _build_index(sdata)
            s_ep_map      = _build_ep_site_map(sdata, sidx)
            s_risk, s_unw = _build_wo_risk_maps(sdata, sidx)
            s_dup         = _build_duplicate_ip_set(sdata)
            s_approved    = _collect_approved_objects(sdata, sidx)
            s_risk, s_unw = _strip_approved_outlets(s_risk, s_unw, _approved_keys(sdata))
            _prepare_report_meta(sdata, sidx, s_risk, s_unw, s_dup, vlan_config)
            s_sites = get_all_sites(sdata)

            _company_heading(doc, comp.get("name", "Bedrijf"), is_first=(i == 0))

            _first = True
            def _rb(title: str, render: bool) -> bool:
                nonlocal _first
                if not render:
                    return False
                _rubriek(doc, title, page_break=not _first)
                _first = False
                return True

            has_switch  = any(d.get("type") == "switch" for d in sdata.get("devices", []))
            has_wifi    = any(e.get("type") == "access_point" for e in sdata.get("endpoints", []))
            has_risk    = sum(len(v) for v in s_risk.values()) > 0
            has_unw     = sum(len(v) for v in s_unw.values()) > 0
            has_eps     = bool(sdata.get("endpoints"))
            has_devinfo = any(d.get("mac") or d.get("serial") or d.get("ip")
                              for d in sdata.get("devices", []))
            has_detail  = any(
                room.get("racks") or room.get("wall_outlets")
                for site in s_sites for room in site.get("rooms", [])
            )

            # 2.1 Samenvatting (altijd)
            _rb("Samenvatting", True)
            _add_site_summary(doc, sdata, sidx, s_risk)

            # 2.2 Aandachtspunten / actieplan (altijd)
            _rb("Aandachtspunten", True)
            _add_action_plan(doc, sdata, sidx, s_risk, s_unw, s_dup)

            # 2.2b Goedgekeurde uitzonderingen (REVIEW-AP)
            if _rb("Goedgekeurde uitzonderingen", bool(s_approved)):
                _add_approved_exceptions(doc, s_approved)

            # 2.3 Switch-overzicht
            if _rb("Switch-overzicht", has_switch):
                _add_switch_overview(doc, sdata, sidx, s_dup)

            # 2.4 Uplink-overzicht
            if _rb("Uplink-overzicht", _has_uplinks(sdata, sidx)):
                _add_uplink_overview(doc, sdata, sidx)

            # 2.5 Wi-Fi overzicht
            if _rb("Wi-Fi overzicht", has_wifi):
                _add_wifi_overview(doc, sdata, sidx)

            # 2.6 Netwerkdetail per ruimte (racks/devices/poorten/wandpunten/verbindingen)
            if _rb("Netwerkdetail", has_detail):
                for site in s_sites:
                    _add_site_header(doc, site["name"], site.get("location", ""))
                    _spacer(doc, 0.4)
                    for room in site.get("rooms", []):
                        if not room.get("racks") and not room.get("wall_outlets"):
                            continue
                        _add_room_header(doc, room["name"], room.get("floor", ""))
                        _spacer(doc, 0.2)
                        _add_rack_overview(doc, room)
                        for rack in room.get("racks", []):
                            if not rack.get("slots"):
                                continue
                            _add_rack_badge(doc, rack["name"], room["name"])
                            _add_device_table(doc, rack, sidx)
                            _spacer(doc, 0.2)
                            _add_port_table(doc, rack, sdata, sidx)
                            _spacer(doc, 0.5)
                    _add_site_outlets(doc, sdata, sidx, site)
                    _add_site_connections(doc, sdata, sidx, site)
                    _add_site_direct_endpoints(doc, sdata, sidx, site)
                    _add_site_vlans(doc, sdata, sidx, site)

            # 2.7 Risico-wandpunten
            if _rb("Risico-wandpunten", has_risk):
                _add_risk_outlets_section(doc, sdata, sidx, s_risk)

            # 2.8 Onverbonden wandpunten
            if _rb("Onverbonden wandpunten", has_unw):
                _add_unwired_outlets_section(doc, sdata, sidx, s_unw)

            # 2.9 Eindapparaten
            if _rb("Eindapparaten", has_eps):
                _add_endpoints_section(doc, sdata, sidx, s_ep_map)

            # 2.10 Device-info
            if _rb("Device-info", has_devinfo):
                _add_device_info_section(doc, sdata, sidx, s_dup)

        # 3. Globale afsluiting: niet-geplaatste objecten + begrippenlijst + revisie
        _add_orphans_section(doc, full_data, g_idx)
        _add_glossary(doc)
        _add_revision_history(doc, version)

        doc.save(filepath)
        return True, ""

    except Exception:
        import traceback
        return False, traceback.format_exc()



# ---------------------------------------------------------------------------
# Standalone test (wordt niet uitgevoerd in de app)
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    from app.helpers import settings_storage
    data = settings_storage.load_network_data()
    ok, err = render_report_docx(data, "/tmp/networkrapport.docx")
    print("OK" if ok else f"FOUT:\n{err}")